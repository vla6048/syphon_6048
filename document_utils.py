from docx.oxml.ns import qn
from docx.shared import Pt
from openpyxl.styles import PatternFill
from num2words import num2words


def _replace_in_paragraph(paragraph, replacements):
    for key, value in replacements.items():
        replacement = str(value)

        while key in paragraph.text:
            start = paragraph.text.find(key)
            end = start + len(key)
            current = 0
            start_run = start_offset = end_run = end_offset = None

            for idx, run in enumerate(paragraph.runs):
                run_end = current + len(run.text)

                if start_run is None and current <= start < run_end:
                    start_run = idx
                    start_offset = start - current

                if current < end <= run_end:
                    end_run = idx
                    end_offset = end - current
                    break

                current = run_end

            if start_run is None or end_run is None:
                break

            runs = paragraph.runs
            if start_run == end_run:
                run = runs[start_run]
                run.text = run.text[:start_offset] + replacement + run.text[end_offset:]
            else:
                first_run = runs[start_run]
                last_run = runs[end_run]
                first_run.text = first_run.text[:start_offset] + replacement + last_run.text[end_offset:]
                for idx in range(start_run + 1, end_run + 1):
                    runs[idx].text = ""


def replace_text_in_document(doc, replacements):
    for paragraph in doc.paragraphs:
        _replace_in_paragraph(paragraph, replacements)


def replace_in_tables(tables, replacements):
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_paragraph(paragraph, replacements)
                if cell.tables:
                    replace_in_tables(cell.tables, replacements)


def formatting_text(document):
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)


def _remove_xml_children(element, child_tags):
    if element is None:
        return

    for child_tag in child_tags:
        for child in list(element.findall(qn(child_tag))):
            element.remove(child)


def _clear_paragraph_highlights(paragraph):
    _remove_xml_children(paragraph._p.pPr, ("w:shd",))

    for run in paragraph.runs:
        run.font.highlight_color = None
        _remove_xml_children(run._r.rPr, ("w:highlight", "w:shd"))


def clear_document_highlights(document):
    for paragraph in document.paragraphs:
        _clear_paragraph_highlights(paragraph)

    for table in document.tables:
        clear_table_highlights(table)


def clear_table_highlights(table):
    for row in table.rows:
        for cell in row.cells:
            _remove_xml_children(cell._tc.tcPr, ("w:shd",))
            for paragraph in cell.paragraphs:
                _clear_paragraph_highlights(paragraph)
            for nested_table in cell.tables:
                clear_table_highlights(nested_table)


def clear_workbook_highlights(workbook):
    empty_fill = PatternFill(fill_type=None)
    for sheet in workbook.worksheets:
        for row in sheet.iter_rows():
            for cell in row:
                cell.fill = empty_fill


def convert_to_currency_words(amount):
    hryvnia_part = int(amount)
    kopiyka_part = int(round((amount - hryvnia_part) * 100))
    hryvnia_words = num2words(hryvnia_part, lang='uk')
    kopiyka_words = num2words(kopiyka_part, lang='uk')
    return f"{hryvnia_words} гривень {kopiyka_words} копійок"


def format_date(date):
    months_ukr = {
        1: 'січня', 2: 'лютого', 3: 'березня', 4: 'квітня', 5: 'травня', 6: 'червня',
        7: 'липня', 8: 'серпня', 9: 'вересня', 10: 'жовтня', 11: 'листопада', 12: 'грудня'
    }
    day = date.strftime("%d")
    month = months_ukr[date.month]
    year = date.strftime("%Y")
    return f"{day} {month} {year} року", month, year, day


def amount_to_time(protocol_amount):
    work_hours = protocol_amount / 1000
    hours = int(work_hours)
    minutes = int((work_hours - hours) * 60)
    return f"{hours} годин {minutes} хвилин"


def create_table(doc, data, headers):
    table = doc.add_table(rows=1, cols=len(headers))

    hdr_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr_cells[idx].text = header

    for row in data:
        row_cells = table.add_row().cells
        for idx, val in enumerate(row):
            row_cells[idx].text = str(val)

    return table


def replace_table_in_document(doc, marker, table):
    for paragraph in doc.paragraphs:
        if marker in paragraph.text:
            paragraph.clear()
            paragraph._element.addnext(table._element)
            break
    return doc
