from docx import Document

template_path = 'static/docs/M-RI_protocol.docx'
doc = Document(template_path)

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            print(cell.tables)
            for tablee in cell.tables:
                for roww in tablee.rows:
                    for celll in roww.cells:
                        print(celll.text)
