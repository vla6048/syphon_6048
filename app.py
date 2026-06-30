from quart import Quart, render_template, request, jsonify,redirect, url_for, send_file, flash, Blueprint
from quart_auth import QuartAuth, basic_auth_required
from docx import Document
from aiohttp import web
import aiomysql
from docx.oxml.ns import qn
from docx.enum.table import WD_ALIGN_VERTICAL
from dotenv import load_dotenv
from datetime import date, datetime
import asyncio
import os
import calendar
import tempfile
from io import BytesIO
from pathlib import Path
import pandas as pd
from db_manager import DatabaseManager
import random
from math import floor
import openpyxl
import uuid
from document_utils import (
    amount_to_time,
    clear_document_highlights,
    clear_workbook_highlights,
    convert_to_currency_words,
    create_table,
    format_date,
    formatting_text,
    replace_in_tables,
    replace_table_in_document,
    replace_text_in_document,
)
from sql_utils import build_placeholders, quote_qualified_identifier

# Загрузка переменных окружения из .env файла
load_dotenv()


async def convert_office_bytes_to_pdf(source_bytes, source_suffix):
    with tempfile.TemporaryDirectory() as temp_dir:
        temp_path = Path(temp_dir)
        source_path = temp_path / f"document{source_suffix}"
        pdf_path = temp_path / "document.pdf"
        profile_path = temp_path / "lo-profile"
        runtime_path = temp_path / "runtime"
        config_path = temp_path / "config"
        cache_path = temp_path / "cache"
        runtime_path.mkdir(mode=0o700)
        config_path.mkdir()
        cache_path.mkdir()
        source_path.write_bytes(source_bytes)

        env = os.environ.copy()
        env["HOME"] = str(temp_path)
        env["XDG_RUNTIME_DIR"] = str(runtime_path)
        env["XDG_CONFIG_HOME"] = str(config_path)
        env["XDG_CACHE_HOME"] = str(cache_path)

        process = await asyncio.create_subprocess_exec(
            "soffice",
            "--headless",
            f"-env:UserInstallation={profile_path.as_uri()}",
            "--convert-to",
            "pdf",
            "--outdir",
            str(temp_path),
            str(source_path),
            stdout=asyncio.subprocess.PIPE,
            stderr=asyncio.subprocess.PIPE,
            env=env,
        )
        stdout, stderr = await process.communicate()

        if process.returncode != 0 or not pdf_path.exists():
            output = (stderr or stdout).decode("utf-8", errors="replace")
            raise RuntimeError(f"LibreOffice PDF conversion failed: {output}")

        return pdf_path.read_bytes()


class MyApp:
    replace_text_in_document = staticmethod(replace_text_in_document)
    replace_in_tables = staticmethod(replace_in_tables)
    formatting_text = staticmethod(formatting_text)
    convert_to_currency_words = staticmethod(convert_to_currency_words)
    format_date = staticmethod(format_date)
    amount_to_time = staticmethod(amount_to_time)
    create_table = staticmethod(create_table)
    replace_table_in_document = staticmethod(replace_table_in_document)

    def __init__(self):
        # Создание экземпляра Quart
        self.app = Quart(__name__)
        QuartAuth(self.app)
        self.app.secret_key = os.getenv('SECRET_KEY') or os.urandom(24)
        self.app.config["QUART_AUTH_BASIC_USERNAME"] = os.getenv('BUSERNAME')
        self.app.config["QUART_AUTH_BASIC_PASSWORD"] = os.getenv('BPASSWD')

        # Настройка подключения к базам данных
        self.local_db = DatabaseManager(
            host=os.getenv('MYSQL_HOST_LOCAL'),
            user=os.getenv('MYSQL_USER_LOCAL'),
            password=os.getenv('MYSQL_PASSWORD_LOCAL'),
            db=os.getenv('MYSQL_DB_LOCAL')
        )

        self.remote_db = DatabaseManager(
            host=os.getenv('MYSQL_HOST_REMOTE'),
            user=os.getenv('MYSQL_USER_REMOTE'),
            password=os.getenv('MYSQL_PASSWORD_REMOTE'),
            db=os.getenv('MYSQL_DB_REMOTE')
        )

        bp = Blueprint('generate_protocols', __name__)
        # Настройка маршрутов
        self.setup_routes()
        self.setup_lifecycle()

    async def __call__(self, scope, receive, send):
        await self.app(scope, receive, send)

    def setup_lifecycle(self):
        @self.app.after_serving
        async def close_database_pools():
            await self.local_db.close()
            await self.remote_db.close()

    def setup_routes(self):

        @self.app.route('/check_payments', methods=['GET', 'POST'])
        @basic_auth_required()
        async def check_payments():
            # Шаг 1: Получение списка инженеров для фильтра (без изменений)
            engineers_query = "SELECT name, inn FROM credentials.ri_credentials;"
            engineers_data = await self.local_db.execute_query(engineers_query)
            engineers = [(row[0], row[1]) for row in engineers_data]

            # Данные для формы (без изменений)
            current_year = date.today().year
            years = list(range(2024, current_year + 2))
            months_data = {
                1: "Январь", 2: "Февраль", 3: "Март", 4: "Апрель", 5: "Май", 6: "Июнь",
                7: "Июль", 8: "Август", 9: "Сентябрь", 10: "Октябрь", 11: "Ноябрь", 12: "Декабрь"
            }

            # Инициализация переменных
            grouped_results = {}
            selected_month = None
            selected_year = None
            selected_engineer_inn = None
            grand_total = 0.0

            if request.method == 'POST':
                form_data = await request.form
                selected_month = form_data.get('month')
                selected_year = form_data.get('year')
                selected_engineer_inn = form_data.get('engineer_inn')

                if not selected_month or not selected_year:
                    await flash("Пожалуйста, выберите месяц и год.", "error")
                    return redirect(url_for('check_payments'))

                # Шаг 3: Обновленный SQL-запрос
                query = """
                        SELECT GROUP_CONCAT(DISTINCT ft.canton SEPARATOR ', ') AS cantons, \
                               agr.agreement_name, \
                               fc.name                                         AS fop_name, \
                               fc.inn                                          AS fop_inn, \
                               rc.name                                         AS ri_name, \
                               rc.inn                                          AS ri_inn, \
                               prot.proto_sum, \
                               agr.id, \
                               agr.agreement_state
                        FROM credentials.agreements agr
                                 JOIN credentials.fop_credentials fc ON agr.master_id = fc.id
                                 JOIN credentials.ri_credentials rc ON agr.ri_id = rc.id
                                 JOIN credentials.protocols prot ON agr.id = prot.agreement
                                 LEFT JOIN credentials.fop_territory ft ON agr.master_id = ft.master_id
                        WHERE YEAR (prot.proto_date) = %s
                          AND MONTH (prot.proto_date) = %s
                          AND prot.proto_state = 1 \
                        """
                params = [selected_year, selected_month]

                if selected_engineer_inn:
                    # ❗ Важно: в этой строке перед AND должен быть пробел, он есть в конце многострочного литерала выше.
                    query += " AND rc.inn = %s"
                    params.append(selected_engineer_inn)

                query += """
        GROUP BY
            agr.agreement_name,
            fc.name, fc.inn,
            rc.name, rc.inn,
            prot.proto_sum,
            agr.id,
            agr.agreement_state
        ORDER BY 
            cantons;
        """

                # Временно выведем запрос для отладки, если ошибка не исчезнет:
                # print("--- SQL Query Start ---")
                # print(query % tuple(params))
                # print("--- SQL Query End ---")

                results_data = await self.local_db.execute_query(query, tuple(params))

                # Шаг 4: Группировка результатов по кантону и подсчет промежуточных сумм
                if results_data:
                    current_canton = None
                    canton_total = 0.0

                    # 9 колонок в результате (cantons, agreement_name, fop_name, fop_inn, ri_name, ri_inn, proto_sum, agr_id, agr_state)
                    for row in results_data:
                        cantons_str = row[0] if row[0] is not None else 'Нет кантона'
                        proto_sum = float(row[6])
                        agr_id = row[7]
                        agr_state = row[8]

                        # Если кантон изменился, сохраняем предыдущий итог
                        if cantons_str != current_canton and current_canton is not None:
                            grouped_results[current_canton]['total'] = canton_total
                            canton_total = 0.0

                        if cantons_str not in grouped_results:
                            grouped_results[cantons_str] = {'records': [], 'total': 0.0}

                        grouped_results[cantons_str]['records'].append({
                            'cantons': cantons_str,
                            'agreement_name': row[1],
                            'fop_name': row[2],
                            'fop_inn': row[3],
                            'ri_name': row[4],
                            'ri_inn': row[5],
                            'proto_sum': proto_sum,
                            'agreement_id': agr_id,
                            'agreement_state': agr_state
                        })

                        canton_total += proto_sum
                        grand_total += proto_sum
                        current_canton = cantons_str

                    # Сохраняем итог для последнего кантона
                    if current_canton is not None:
                        grouped_results[current_canton]['total'] = canton_total

                    await flash(f"Найдено {len(results_data)} записей для сверки.", "success")
                else:
                    await flash("Протоколы для выбранных критериев не найдены.", "info")

            # Шаг 5: Отображение шаблона
            return await render_template(
                'check_payments.html',
                months=months_data.items(),
                years=years,
                engineers=engineers,
                grouped_results=grouped_results,
                grand_total=grand_total,
                selected_month=selected_month,
                selected_year=selected_year,
                selected_engineer_inn=selected_engineer_inn
            )

        @self.app.route('/get_unified_values/<related_table>', methods=['GET'])
        @basic_auth_required()
        async def get_unified_values(related_table):
            """
            Возвращает унифицированные значения для выбранного типа оборудования.
            """
            query_unified = """
                SELECT category, value 
                FROM equipments.unified_values 
                WHERE related_table = %s
            """
            unified_values_raw = await self.local_db.execute_query(query_unified, (related_table,))

            # Группируем по категориям
            unified_values = {}
            for category, value in unified_values_raw:
                unified_values.setdefault(category, []).append(value)

            return jsonify(unified_values)

        @self.app.route('/equipment_insertion', methods=['GET', 'POST'])
        @basic_auth_required()
        async def equipment_insertion():
            # Получаем список типов оборудования
            query_types = "SELECT id, name, related_table FROM equipments.equipment_types"
            equipment_types = await self.local_db.execute_query(query_types)

            # Преобразуем в список словарей для удобства использования
            equipment_types_dict = {str(t[0]): {"name": t[1], "related_table": t[2]} for t in equipment_types}

            if request.method == 'POST':
                data = await request.form
                type_id = data.get('type_id')
                state = data.get('state')
                remark = data.get('remark') or None
                sn = data.get('sn') if data.get('sn') else f'NO_SN_{uuid.uuid4().hex}'

                if not type_id or type_id not in equipment_types_dict:
                    return jsonify({"success": False, "error": "Некорректный тип оборудования"}), 400

                related_table = equipment_types_dict[type_id]["related_table"]

                # Вставка в таблицу equipment
                insert_equipment = """
                    INSERT INTO equipments.equipment (type_id, state, sn, remark)
                    VALUES (%s, %s, %s, %s)
                """
                equipment_values = (type_id, state, sn, remark)
                equipment_id = await self.local_db.execute_insert(insert_equipment, equipment_values)

                # Вставка в связанную таблицу (если она есть)
                if related_table and equipment_id:
                    try:
                        quoted_related_table = quote_qualified_identifier(f"equipments.{related_table}")
                    except ValueError:
                        return jsonify({"success": False, "error": "Некорректная связанная таблица"}), 400

                    query_columns = f"SHOW COLUMNS FROM {quoted_related_table}"
                    columns_data = await self.local_db.execute_query(query_columns)
                    columns = [col[0] for col in columns_data if col[0] not in ('id', 'equipment_id')]

                    related_values = [equipment_id] + [data.get(col) for col in columns]
                    related_columns = ["equipment_id"] + columns
                    quoted_columns = ", ".join(quote_qualified_identifier(col) for col in related_columns)
                    insert_related = f"""
                        INSERT INTO {quoted_related_table} ({quoted_columns})
                        VALUES ({', '.join(['%s'] * len(related_values))})
                    """
                    await self.local_db.execute_insert(insert_related, related_values)

                return jsonify({"success": True, "equipment_id": equipment_id})

            return await render_template('equipment_insertion.html', equipment_types=equipment_types)

        @self.app.route('/check_sn', methods=['GET'])
        @basic_auth_required()
        async def check_sn():
            sn = request.args.get('sn')
            if not sn:
                return jsonify({"exists": False})

            query = "SELECT COUNT(*) FROM equipments.equipment WHERE sn = %s"
            result = await self.local_db.execute_query(query, (sn,))

            return jsonify({"exists": result[0][0] > 0})

        @self.app.route('/get_fields/<table>', methods=['GET'])
        async def get_fields(table):
            """ Возвращает список полей таблицы (кроме 'id' и 'equipment_id') в формате JSON. """
            try:
                quoted_table = quote_qualified_identifier(f"equipments.{table}")
            except ValueError:
                return jsonify({"error": "Некорректное имя таблицы"}), 400

            query = f"SHOW COLUMNS FROM {quoted_table}"
            result = await self.local_db.execute_query(query)

            if not result:
                return jsonify([])

            # Исключаем технические столбцы
            columns = [col[0] for col in result if col[0] not in ("id", "equipment_id")]
            return jsonify(columns)

        @self.app.route('/llc_acts/<int:act_id>/generate_report_llc', methods=['POST'])
        async def generate_report_llc(act_id):
            # SQL-запрос для получения данных по act_id
            query = """
                    SELECT act.act_date, 
                           act.act_sum, 
                           act.id, 
                           agr.agreement_name, 
                           agr.agreement_date, 
                           llc.name, 
                           llc.edrpou, 
                           ri.name, 
                           ri.iban, 
                           ri.bank_account_detail, 
                           ri.address, 
                           ri.phone, 
                           ri.inn, 
                           ri.name_short,
                           llc.in_persona,
                           ri.pidstava,
                           llc.address,
                           llc.iban,
                           llc.bank_account_detail,
                           llc.inn,
                           llc.name_short
                    FROM credentials.llc_acts act
                    JOIN credentials.llc_agreements agr ON act.agreement = agr.id
                    JOIN credentials.llc_credentials llc ON agr.llc_id = llc.id
                    JOIN credentials.ri_credentials ri ON agr.ri_id = ri.id
                    WHERE act.id = %s
                    """
            result = await self.local_db.execute_query(query, (act_id,))
            data = result[0] if result else {}

            # Новый запрос для получения данных из llc_acts_data
            query_acts_data = """
                    SELECT sw_rank, model_list, count_devices, ip_list, worktime_float
                    FROM credentials.llc_acts_data
                    WHERE act_id = %s
                    """
            acts_data_result = await self.local_db.execute_query(query_acts_data, (act_id,))

            # Подготовка данных для замены
            replacements = {
                "@act_name": f"R{act_id}_{data[0].strftime('%m/%y')}_{data[3]}",
                "@act_date": f"«{calendar.monthrange(int(data[0].strftime('%Y')), int(data[0].month))[-1]}» {self.format_date(data[0])[1]} {self.format_date(data[0])[2]} року",
                "@ri_name": data[7],
                "@ri_iban": data[8],
                "@bank_account_detail_ri": data[9],
                "@ri_address": data[10],
                "@ri_phone": data[11],
                "@ri_inn": data[12],
                "@llc_name": data[5],
                "@llc_edrpou": data[6],
                "@agr_name": data[3],
                "@agr_date": self.format_date(data[4])[0],
                "@act_sum": data[1],
                "@actsumwords": self.convert_to_currency_words(data[1]),  # Конвертация числа в пропись
                "@ri_shortname": data[13],
                "@llc_in_persona": data[14],
                "@ri_pidstava": data[15],
                "@current_month": self.format_date(data[0])[1],
                "@current_year": self.format_date(data[0])[2],
                "@last_day_of_the_month": calendar.monthrange(int(data[0].strftime('%Y')), int(data[0].month))[-1],
                "@llc_address": data[16],
                "@llc_iban": data[17],
                "@bank_account_detail_llc": data[18],
                "@llc_inn": data[19],
                "@llc_shortname": data[20]
            }

            # Условия для замены значений из acts_data
            if str(data[6]) == '38736443':  # Если edrpou == 38736443
                for row in acts_data_result:
                    if row[0] == 1:  # sw_rank == 1
                        replacements["@rank1_models"] = row[1]
                        replacements["@rank1_count"] = row[2]
                        replacements["@rank1_ips"] = row[3]
                    elif row[0] == 2:  # sw_rank == 2
                        replacements["@rank2_models"] = row[1]
                        replacements["@rank2_count"] = row[2]
                        replacements["@rank2_ips"] = row[3]
                    elif row[0] == 0:  # sw_rank == 0
                        replacements["@time_models"] = row[1]
                        replacements["@time_count"] = round(row[4], 2)
                        replacements["@time_ips"] = row[3]
            else:  # Если edrpou != 38736443
                for row in acts_data_result:
                    if row[0] == 4:  # sw_rank == 4
                        replacements["@rank4_models"] = row[1]
                        replacements["@rank4_count"] = row[2]
                        replacements["@rank4_ips"] = row[3]
                    elif row[0] == 3:  # sw_rank == 3
                        replacements["@rank3_models"] = row[1]
                        replacements["@rank3_count"] = row[2]
                        replacements["@rank3_ips"] = row[3]
                    elif row[0] == 0:  # sw_rank == 0
                        replacements["@time_models"] = row[1]
                        replacements["@time_count"] = round(row[4], 2)
                        replacements["@time_ips"] = row[3]

            # Проверка `llc_edrpou` и выбор пути к шаблону
            template_path = 'static/docs/kdn_report.docx' if str(data[6]) == '38736443' else 'static/docs/llc_report.docx'

            # Открываем шаблон документа Word
            document = Document(template_path)

            # Замена меток на соответствующие значения в тексте документа
            self.replace_text_in_document(document, replacements)
            self.replace_in_tables(document.tables, replacements)
            self.formatting_text(document)

            # Сохраняем измененный файл в памяти
            output = BytesIO()
            document.save(output)
            output.seek(0)
            docx_name = f'{data[3]}_Звіт_{self.format_date(data[0])[1]}_{self.format_date(data[0])[2]}'

            # Отправляем файл для скачивания
            return await send_file(output, as_attachment=True, attachment_filename=f"{docx_name}.docx")

        @self.app.route('/llc_acts/<int:act_id>/generate_act', methods=['POST'])
        async def generate_act(act_id):
            # SQL-запрос для получения данных по act_id
            query = """
            SELECT act.act_date, 
                   act.act_sum, 
                   act.id, 
                   agr.agreement_name, 
                   agr.agreement_date, 
                   llc.name, 
                   llc.edrpou, 
                   ri.name, 
                   ri.iban, 
                   ri.bank_account_detail, 
                   ri.address, 
                   ri.phone, 
                   ri.inn, 
                   ri.name_short,
                   llc.in_persona,
                   ri.pidstava,
                   llc.address,
                   llc.iban,
                   llc.bank_account_detail,
                   llc.inn,
                   llc.name_short
            FROM credentials.llc_acts act
            JOIN credentials.llc_agreements agr ON act.agreement = agr.id
            JOIN credentials.llc_credentials llc ON agr.llc_id = llc.id
            JOIN credentials.ri_credentials ri ON agr.ri_id = ri.id
            WHERE act.id = %s
            """
            result = await self.local_db.execute_query(query, (act_id,))
            data = result[0] if result else {}

            # Новый запрос для получения данных из llc_acts_data
            query_acts_data = """
                    SELECT sw_rank, count_devices, worktime_float
                    FROM credentials.llc_acts_data
                    WHERE act_id = %s
                    """
            acts_data_result = await self.local_db.execute_query(query_acts_data, (act_id,))


            # Подготовка данных для замены
            replacements = {
                "@act_name": f"A{act_id}_{data[0].strftime('%m/%y')}_{data[3]}",
                "@act_date": f"«{calendar.monthrange(int(data[0].strftime('%Y')), int(data[0].month))[-1]}» {self.format_date(data[0])[1]} {self.format_date(data[0])[2]} року",
                "@ri_name": data[7],
                "@ri_iban": data[8],
                "@bank_account_detail_ri": data[9],
                "@ri_address": data[10],
                "@ri_phone": data[11],
                "@ri_inn": data[12],
                "@llc_name": data[5],
                "@llc_edrpou": data[6],
                "@agr_name": data[3],
                "@agr_date": self.format_date(data[4])[0],
                "@act_sum": data[1],
                "@actsumwords": self.convert_to_currency_words(data[1]),  # Конвертация числа в пропись
                "@ri_shortname": data[13],
                "@llc_in_persona": data[14],
                "@ri_pidstava": data[15],
                "@current_month": self.format_date(data[0])[1],
                "@current_year": self.format_date(data[0])[2],
                "@last_day_of_the_month": calendar.monthrange(int(data[0].strftime('%Y')), int(data[0].month))[-1],
                "@llc_address": data[16],
                "@llc_iban": data[17],
                "@bank_account_detail_llc": data[18],
                "@llc_inn": data[19],
                "@llc_shortname": data[20]
            }

            # Условия для замены значений из acts_data
            if str(data[6]) == '38736443':  # Если edrpou == 38736443
                for row in acts_data_result:
                    if row[0] == 1:  # sw_rank == 1
                        replacements["@rank1_count"] = row[1]
                        replacements["@rank1_sum"] = float(row[1] * 1000)
                        sum_rank1 = float(row[1] * 1000)
                    elif row[0] == 2:  # sw_rank == 2
                        replacements["@rank2_count"] = row[1]
                        replacements["@rank2_sum"] = float(row[1] * 1000)
                        sum_rank2 = float(row[1] * 1000)
                    elif row[0] == 0:  # sw_rank == 0
                        replacements["@time_count"] = round(row[2], 2)
                        replacements["@time_sum"] = round(float(data[1]) - sum_rank2 - sum_rank1, 2)
            else:  # Если edrpou != 38736443
                for row in acts_data_result:
                    if row[0] == 4:  # sw_rank == 4
                        replacements["@rank4_count"] = row[1]
                        replacements["@rank4_sum"] = float(row[1] * 500)
                        sum_rank4 = float(row[1] * 500)
                    elif row[0] == 3:  # sw_rank == 3
                        replacements["@rank3_count"] = row[1]
                        replacements["@rank3_sum"] = float(row[1] * 1000)
                        sum_rank3 = float(row[1] * 1000)
                    elif row[0] == 0:  # sw_rank == 0
                        replacements["@time_count"] = round(row[2], 2)
                        replacements["@time_sum"] = round(float(data[1]) - sum_rank4 - sum_rank3, 2)


            # Проверка `llc_edrpou` и выбор пути к шаблону
            template_path = 'static/docs/kdn_act.docx' if str(data[6]) == '38736443' else 'static/docs/llc_act.docx'

            # Открываем шаблон документа Word
            document = Document(template_path)

            # Замена меток на соответствующие значения в тексте документа
            self.replace_text_in_document(document, replacements)
            self.replace_in_tables(document.tables, replacements)
            self.formatting_text(document)

            # Сохраняем измененный файл в памяти
            output = BytesIO()
            document.save(output)
            output.seek(0)
            docx_name = f'{data[3]}_Акт_{self.format_date(data[0])[1]}_{self.format_date(data[0])[2]}'

            # Отправляем файл для скачивания
            return await send_file(output, as_attachment=True, attachment_filename=f"{docx_name}.docx",
                                   mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        @self.app.route('/llc_acts/<int:act_id>/generate_bill', methods=['POST'])
        async def generate_bill(act_id):
            # SQL-запрос для получения данных по act_id
            query = """
            SELECT act.act_date, 
                   act.act_sum, 
                   act.id, 
                   agr.agreement_name, 
                   agr.agreement_date, 
                   llc.name, 
                   llc.edrpou, 
                   ri.name, 
                   ri.iban, 
                   ri.bank_account_detail, 
                   ri.address, 
                   ri.phone, 
                   ri.inn, 
                   ri.name_short
            FROM credentials.llc_acts act
            JOIN credentials.llc_agreements agr ON act.agreement = agr.id
            JOIN credentials.llc_credentials llc ON agr.llc_id = llc.id
            JOIN credentials.ri_credentials ri ON agr.ri_id = ri.id
            WHERE act.id = %s
            """
            result = await self.local_db.execute_query(query, (act_id,))
            data = result[0] if result else {}

            # Новый запрос для получения данных из llc_acts_data
            query_acts_data = """
            SELECT sw_rank, count_devices, worktime_float
            FROM credentials.llc_acts_data
            WHERE act_id = %s
            """
            acts_data_result = await self.local_db.execute_query(query_acts_data, (act_id,))

            # Подготовка переменных для replacement
            replacements = {
                "@bill_name": f"B{act_id}_{data[0].strftime('%m/%y')}_{data[3]}",
                "@bill_date": f"{calendar.monthrange(int(data[0].strftime('%Y')), int(data[0].month))[-1]} {self.format_date(data[0])[1]} {self.format_date(data[0])[2]} року",
                "@ri_name": data[7],
                "@ri_iban": data[8],
                "@bank_account_detail_ri": data[9],
                "@ri_address": data[10],
                "@ri_phone": data[11],
                "@ri_inn": data[12],
                "@llc_name": data[5],
                "@llc_edrpou": data[6],
                "@agr_name": data[3],
                "@agr_date": self.format_date(data[4])[0],
                "@bill_sum": data[1],
                "@handwritebill_sum": self.convert_to_currency_words(data[1]),  # Конвертация числа в пропись
                "@ri_shortname": data[13]
            }

            # Условия для замены значений из acts_data
            if str(data[6]) == '38736443':  # Если edrpou == 38736443
                for row in acts_data_result:
                    if row[0] == 1:  # sw_rank == 1
                        replacements["@rank1_count"] = row[1]
                        replacements["@rank1_sum"] = float(row[1] * 1000)
                        sum_rank1 = float(row[1] * 1000)
                    elif row[0] == 2:  # sw_rank == 2
                        replacements["@rank2_count"] = row[1]
                        replacements["@rank2_sum"] = float(row[1] * 1000)
                        sum_rank2 = float(row[1] * 1000)
                    elif row[0] == 0:  # sw_rank == 0
                        replacements["@time_count"] = round(row[2], 2)
                        replacements["@time_sum"] = round(float(data[1]) - sum_rank2 - sum_rank1, 2)
            else:  # Если edrpou != 38736443
                for row in acts_data_result:
                    if row[0] == 4:  # sw_rank == 4
                        replacements["@rank4_count"] = row[1]
                        replacements["@rank4_sum"] = float(row[1]*500)
                        sum_rank4 = float(row[1]*500)
                    elif row[0] == 3:  # sw_rank == 3
                        replacements["@rank3_count"] = row[1]
                        replacements["@rank3_sum"] = float(row[1]*1000)
                        sum_rank3 = float(row[1]*1000)
                    elif row[0] == 0:  # sw_rank == 0
                        replacements["@time_count"] = round(row[2],2)
                        replacements["@time_sum"] = round(float(data[1])-sum_rank4-sum_rank3, 2)



            # Открываем файл шаблона
            template_path = 'static/docs/kdn_bill.xlsx' if str(data[6]) == '38736443' else 'static/docs/llc_bill.xlsx'
            workbook = openpyxl.load_workbook(template_path)
            sheet = workbook.active

            # Замена меток на соответствующие значения
            for row in sheet.iter_rows():
                for cell in row:
                    if isinstance(cell.value, str):
                        # Проверяем наличие и заменяем все метки внутри текста ячейки
                        for key, replacement in replacements.items():
                            replacement = str(replacement)
                            if key in cell.value:
                                cell.value = cell.value.replace(key, replacement)

            # Сохраняем измененный файл в памяти
            output = BytesIO()
            workbook.save(output)
            output.seek(0)
            xlxs_name = f'{data[3]}_Рахунок_{self.format_date(data[0])[1]}_{self.format_date(data[0])[2]}'

            # Отправляем файл для скачивания
            return await send_file(output, as_attachment=True, attachment_filename=f"{xlxs_name}.xlsx",
                                   mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

        @self.app.route('/llc_acts/<int:agreement_id>/generate_data/<int:act_id>', methods=['POST'])
        # @basic_auth_required()
        async def generate_act_data(agreement_id, act_id):
            # Шаг 1. Получение данных из таблицы llc_acts
            act_query = """
            SELECT act_date, act_sum, agreement
            FROM credentials.llc_acts
            WHERE id = %s;
            """
            act_data = await self.local_db.execute_query(act_query, (act_id,))

            if not act_data:
                return "Акт не найден", 404

            act_date, act_sum, agreement = act_data[0]

            # Шаг 2. Получение llc_id и ri_id по agreement
            agreement_query = """
            SELECT llc_id, ri_id
            FROM credentials.llc_agreements
            WHERE id = %s;
            """
            agreement_data = await self.local_db.execute_query(agreement_query, (agreement,))

            if not agreement_data:
                return "Договор не найден", 404

            llc_id, ri_id = agreement_data[0]

            # Шаг 3. Получение edrpou по llc_id
            edrpou_query = """
            SELECT edrpou
            FROM credentials.llc_credentials
            WHERE id = %s;
            """
            edrpou_data = await self.local_db.execute_query(edrpou_query, (llc_id,))

            if not edrpou_data:
                return "Организация не найдена", 404

            edrpou = edrpou_data[0][0]

            # Шаг 4. Логика в зависимости от edrpou
            if edrpou == 38736443:
                # Первая логика
                await handle_kdn_logic(act_date, act_sum, agreement, ri_id, act_id)
            else:
                # Вторая логика
                await handle_llc_logic(act_sum, llc_id, act_id)

            return redirect(url_for('llc_acts', agreement_id=agreement_id))

        async def handle_llc_logic(act_sum, llc_id, act_id, data_table="credentials.llc_acts_data"):
            quoted_data_table = quote_qualified_identifier(data_table)
            # Установленные цены за единицу для каждого ранга
            RANK3_COST = 1000
            RANK4_COST = 500
            CONSULTATION_COST = 1000

            # Генерируем случайный процент для ранга 3 и ранга 4
            rank3_percentage = random.uniform(0.6, 0.7)  # 60% - 70%
            rank4_percentage = random.uniform(0.15, 0.2)  # 15% - 20%
            consultation_percentage = 1.0 - rank3_percentage - rank4_percentage  # Остаток на консультации

            # Рассчитываем бюджет для каждого ранга
            rank3_budget = act_sum * rank3_percentage
            rank4_budget = act_sum * rank4_percentage
            consultation_budget = act_sum * consultation_percentage

            # Получение данных оборудования ранга 3
            rank3_query = """
            SELECT sw.id, sw.model, sw.ip
            FROM dbsyphon.switches_report sw
            JOIN credentials.llc_cantons cant ON sw.canton = cant.canton
            WHERE sw.switch_rank = 3 AND cant.llc_id = %s;
            """
            rank3_data = await self.local_db.execute_query(rank3_query, (llc_id,))
            rank3_models, rank3_ips = zip(*[(row[1], row[2]) for row in rank3_data]) if rank3_data else ([], [])

            # Получение данных оборудования ранга 4
            rank4_query = """
            SELECT sw.id, sw.model, sw.ip
            FROM dbsyphon.switches_report sw
            JOIN credentials.llc_cantons cant ON sw.canton = cant.canton
            WHERE sw.switch_rank = 4 AND cant.llc_id = %s;
            """
            rank4_data = await self.local_db.execute_query(rank4_query, (llc_id,))
            rank4_models, rank4_ips = zip(*[(row[1], row[2]) for row in rank4_data]) if rank4_data else ([], [])

            # Расчет максимального количества устройств для ранга 3
            rank3_units = min(floor(rank3_budget / RANK3_COST), len(rank3_models))
            remaining_rank3_budget = rank3_budget - (rank3_units * RANK3_COST)
            selected_rank3 = list(zip(rank3_models, rank3_ips))[:rank3_units]

            # Расчет максимального количества устройств для ранга 4
            rank4_units = min(floor(rank4_budget / RANK4_COST), len(rank4_models))
            remaining_rank4_budget = rank4_budget - (rank4_units * RANK4_COST)
            selected_rank4 = list(zip(rank4_models, rank4_ips))[:rank4_units]

            # Если остались средства после оборудования, они пойдут на консультации
            remaining_budget = remaining_rank3_budget + remaining_rank4_budget + consultation_budget
            consultation_hours = floor(remaining_budget / CONSULTATION_COST)
            consultation_minutes = round((remaining_budget % CONSULTATION_COST) / CONSULTATION_COST * 60)
            consultation_time_in_float = consultation_hours + (
                        consultation_minutes / 60.0) if remaining_budget > 0 else 0.0
            total_consultation_minutes = consultation_hours * 60 + consultation_minutes

            # Расчет количества устройств для консультаций: одно устройство за каждые 10 минут
            total_consultation_devices = total_consultation_minutes // 10


            # Отбор оборудования для консультаций (миксуем оборудование рангов 3 и 4)
            all_models = rank3_models + rank4_models
            all_ips = rank3_ips + rank4_ips
            selected_consultation = list(zip(all_models, all_ips))[:total_consultation_devices]


            # Формирование отчета для ранга 3
            rank3_report = {
                "models": [x[0] for x in selected_rank3],
                "ips": [x[1] for x in selected_rank3],
                "count": rank3_units
            }

            # Формирование отчета для ранга 4
            rank4_report = {
                "models": [x[0] for x in selected_rank4],
                "ips": [x[1] for x in selected_rank4],
                "count": rank4_units
            }

            # Формирование отчета для консультаций
            consultation_report = {
                "models": [x[0] for x in selected_consultation],
                "ips": [x[1] for x in selected_consultation],
                "time": f"{consultation_time_in_float:.2f} ч"
            }

            # Запись данных в таблицу `credentials.llc_acts_data`
            # Ранг 3
            await self.local_db.execute_query(f"""
                INSERT INTO {quoted_data_table}
                (act_id, sw_rank, model_list, count_devices, ip_list, worktime_float)
                VALUES (%s, %s, %s, %s, %s, %s)
            """, (
                act_id, 3, '\n'.join(rank3_report["models"]), rank3_report["count"], '\n'.join(rank3_report["ips"]), 0))

            # Ранг 4
            await self.local_db.execute_query(f"""
                INSERT INTO {quoted_data_table}
                (act_id, sw_rank, model_list, count_devices, ip_list, worktime_float)
                VALUES (%s, %s, %s, %s, %s, %s)
            """, (
                act_id, 4, '\n'.join(rank4_report["models"]), rank4_report["count"], '\n'.join(rank4_report["ips"]), 0))

            # Консультация
            if consultation_time_in_float > 0:
                await self.local_db.execute_query(f"""
                    INSERT INTO {quoted_data_table}
                    (act_id, sw_rank, model_list, count_devices, ip_list, worktime_float)
                    VALUES (%s, %s, %s, %s, %s, %s)
                """, (act_id, 0, '\n'.join(consultation_report["models"]), total_consultation_devices,
                      '\n'.join(consultation_report["ips"]), consultation_time_in_float))

            # Отладочная информация
            print(f"Настройка оборудования ранга 3 {str(rank3_report)}")
            print(f"Настройка оборудования ранга 4 {str(rank4_report)}")
            print(f"Консультация по работе оборудования {str(consultation_report)}")


        async def handle_kdn_logic(
                act_date,
                act_sum,
                agreement,
                ri_id,
                act_id,
                data_table="credentials.llc_acts_data",
                cantons_table="credentials.engineer_cantons"
        ):
            quoted_data_table = quote_qualified_identifier(data_table)
            quoted_cantons_table = quote_qualified_identifier(cantons_table)
            # Установленная цена за единицу оборудования или времени консультаций
            UNIT_COST = 1000

            # Процентное распределение бюджета
            rank1_share = random.uniform(0.10, 0.40)
            rank2_share = random.uniform(0.40, 0.60)
            consultation_share = 1 - (rank1_share + rank2_share)

            # Расчет бюджета для каждого ранга и консультаций
            rank1_budget = act_sum * rank1_share
            rank2_budget = act_sum * rank2_share
            consultation_budget = act_sum * consultation_share

            # Получение данных оборудования ранга 1
            rank1_query = f"""
            SELECT sw.model, sw.ip
            FROM dbsyphon.switches_report sw
            JOIN {quoted_cantons_table} cant ON sw.canton = cant.canton
            WHERE sw.switch_rank = 1 AND cant.engineer_id = %s;
            """
            rank1_data = await self.local_db.execute_query(rank1_query, (ri_id,))

            # Получение данных оборудования ранга 2
            rank2_query = f"""
            SELECT sw.model, sw.ip
            FROM dbsyphon.switches_report sw
            JOIN {quoted_cantons_table} cant ON sw.canton = cant.canton
            WHERE sw.switch_rank = 2 AND cant.engineer_id = %s;
            """
            rank2_data = await self.local_db.execute_query(rank2_query, (ri_id,))

            # Вычисление количества необходимых устройств для каждого ранга
            rank1_units_target = floor(rank1_budget / UNIT_COST)
            rank2_units_target = floor(rank2_budget / UNIT_COST)

            selected_rank1 = random.sample(rank1_data, min(rank1_units_target, len(rank1_data))) if rank1_data else []
            selected_rank2 = random.sample(rank2_data, min(rank2_units_target, len(rank2_data))) if rank2_data else []
            rank1_units = len(selected_rank1)
            rank2_units = len(selected_rank2)
            selected_rank1_models = [row[0] for row in selected_rank1]
            selected_rank1_ips = [row[1] for row in selected_rank1]
            selected_rank2_models = [row[0] for row in selected_rank2]
            selected_rank2_ips = [row[1] for row in selected_rank2]

            # Расчет времени консультаций, если осталась сумма
            remaining_budget = (rank1_budget - rank1_units * UNIT_COST) + \
                               (rank2_budget - rank2_units * UNIT_COST) + \
                               consultation_budget
            consultation_time_in_float = round(remaining_budget / UNIT_COST, 2)

            # Формирование отчета для ранга 1
            rank1_report = {
                "models": selected_rank1_models,
                "ips": selected_rank1_ips,
                "count": rank1_units
            }

            # Формирование отчета для ранга 2
            rank2_report = {
                "models": selected_rank2_models,
                "ips": selected_rank2_ips,
                "count": rank2_units
            }

            # Формирование отчета для консультаций
            consultation_report = {
                "models": selected_rank1_models + selected_rank2_models,
                "ips": selected_rank1_ips + selected_rank2_ips,
                "time": f"{consultation_time_in_float:.2f} ч"
            }

            # Запись данных в таблицу `credentials.llc_acts_data`
            # Ранг 1
            await self.local_db.execute_query(f"""
                INSERT INTO {quoted_data_table}
                (act_id, sw_rank, model_list, count_devices, ip_list, worktime_float)
                VALUES (%s, %s, %s, %s, %s, %s)
            """, (
                act_id, 1, '\n'.join(rank1_report["models"]), rank1_report["count"], '\n'.join(rank1_report["ips"]), 0))

            # Ранг 2
            await self.local_db.execute_query(f"""
                INSERT INTO {quoted_data_table}
                (act_id, sw_rank, model_list, count_devices, ip_list, worktime_float)
                VALUES (%s, %s, %s, %s, %s, %s)
            """, (
                act_id, 2, '\n'.join(rank2_report["models"]), rank2_report["count"], '\n'.join(rank2_report["ips"]), 0))

            # Консультации
            if consultation_time_in_float > 0:
                await self.local_db.execute_query(f"""
                    INSERT INTO {quoted_data_table}
                    (act_id, sw_rank, model_list, count_devices, ip_list, worktime_float)
                    VALUES (%s, %s, %s, %s, %s, %s)
                """, (act_id, 0, '\n'.join(consultation_report["models"]), rank1_units + rank2_units,
                      '\n'.join(consultation_report["ips"]),
                      consultation_time_in_float))

            # Отладочная информация
            print(f"Настройка оборудования ранга 1 {str(rank1_report)}")
            print(f"Настройка оборудования ранга 2 {str(rank2_report)}")
            print(f"Консультация по работе оборудования {str(consultation_report)}")

        @self.app.route('/llc_acts/<int:agreement_id>/generate_protocol', methods=['GET'])
        @basic_auth_required()
        async def generate_llc_protocol(agreement_id):
            # Получаем данные по договору и связанные данные для замены в шаблоне
            agreement_query = """
            SELECT la.id AS proto_num, la.agreement_name, la.agreement_date, 
                   lc.name AS llc_name, lc.in_persona, lc.address AS llc_address, lc.edrpou AS llc_edrpou, 
                   lc.iban AS llc_iban, lc.bank_account_detail AS bank_account_detail_llc, lc.inn AS llc_inn, lc.name_short AS llc_shortname,
                   ri.name AS ri_name, ri.inn AS ri_inn, ri.pidstava, ri.address AS ri_address, 
                   ri.iban AS ri_iban, ri.bank_account_detail AS bank_account_detail_ri, ri.name_short AS ri_shortname
            FROM credentials.llc_agreements AS la
            JOIN credentials.llc_credentials AS lc ON la.llc_id = lc.id
            JOIN credentials.ri_credentials AS ri ON la.ri_id = ri.id
            WHERE la.id = %s;
            """
            agreement_data = await self.local_db.execute_query(agreement_query, (agreement_id,))

            # Проверка, что данные по договору найдены
            if not agreement_data:
                return "Договор не найден", 404

            # Извлечение данных из результата
            agreement = agreement_data[0]

            # Проверка ЕДРПОУ организации
            llc_edrpou = agreement[6]
            print(llc_edrpou)
            print(type(llc_edrpou))
            if llc_edrpou == 38736443:
                template_path = 'static/docs/KDN_proto.docx'
            else:
                # Загружаем шаблон llc_proto.docx
                template_path = 'static/docs/llc_proto.docx'

            doc = Document(template_path)
            # Формируем номер протокола
            proto_num = f"{agreement[0]}_{agreement[2].strftime('%Y-%m-%d')}_{agreement[1]}"

            # Задаем значения для замены
            replacements = {
                '@proto_num': proto_num,
                '@agr_name': agreement[1],
                '@agr_date': self.format_date(agreement[2])[0],
                '@llc_name': agreement[3],
                '@persona': agreement[4],
                '@ri_name': agreement[11],
                '@ri_inn': agreement[12],
                '@pidstava': agreement[13],
                '@llc_address': agreement[5],
                '@llc_edrpou': llc_edrpou,
                '@llc_iban': agreement[7],
                '@bank_account_detail_llc': agreement[8],
                '@llc_inn': agreement[9],
                '@llc_shortname': agreement[10],
                '@ri_address': agreement[14],
                '@ri_iban': agreement[15],
                '@bank_account_detail_ri': agreement[16],
                '@ri_shortname': agreement[17]
            }

            # Замена текста в шаблоне
            self.replace_text_in_document(doc, replacements)
            self.replace_in_tables(doc.tables, replacements)

            # Сохранение документа
            doc_io = BytesIO()
            doc.save(doc_io)
            doc_io.seek(0)

            # Формируем название файла
            file_name = f"{agreement[1]} Протокол.docx"

            # Отправка документа клиенту
            return await send_file(doc_io, as_attachment=True, attachment_filename=file_name)

        @self.app.route('/llc_acts/<int:agreement_id>/generate_contract', methods=['GET'])
        @basic_auth_required()
        async def generate_llc_contract(agreement_id):
            # Получаем данные по договору и связанные данные для замены в шаблоне
            agreement_query = """
            SELECT la.id AS contract_num, la.agreement_name, la.agreement_date, 
                   lc.name AS llc_name, lc.in_persona, lc.address AS llc_address, lc.edrpou AS llc_edrpou, 
                   lc.iban AS llc_iban, lc.bank_account_detail AS bank_account_detail_llc, lc.inn AS llc_inn, lc.name_short AS llc_shortname,
                   ri.name AS ri_name, ri.inn AS ri_inn, ri.pidstava, ri.address AS ri_address, 
                   ri.iban AS ri_iban, ri.bank_account_detail AS bank_account_detail_ri, ri.name_short AS ri_shortname
            FROM credentials.llc_agreements AS la
            JOIN credentials.llc_credentials AS lc ON la.llc_id = lc.id
            JOIN credentials.ri_credentials AS ri ON la.ri_id = ri.id
            WHERE la.id = %s;
            """
            agreement_data = await self.local_db.execute_query(agreement_query, (agreement_id,))

            # Проверяем, что данные по договору найдены
            if not agreement_data:
                return "Договор не найден", 404

            # Извлекаем данные из результата
            agreement = agreement_data[0]

            # Определяем шаблон договора
            template_path = 'static/docs/llc_contract.docx'

            # Загружаем шаблон
            doc = Document(template_path)

            # Формируем номер договора
            contract_num = f"{agreement[0]}_{agreement[2].strftime('%Y-%m-%d')}_{agreement[1]}"

            # Задаем значения для замены
            replacements = {
                '@contract_num': contract_num,
                '@agr_name': agreement[1],
                '@agr_date': self.format_date(agreement[2])[0],
                '@llc_name': agreement[3],
                '@persona': agreement[4],
                '@ri_name': agreement[11],
                '@ri_inn': agreement[12],
                '@pidstava': agreement[13],
                '@llc_address': agreement[5],
                '@llc_edrpou': agreement[6],
                '@llc_iban': agreement[7],
                '@bank_account_detail_llc': agreement[8],
                '@llc_inn': agreement[9],
                '@llc_shortname': agreement[10],
                '@ri_address': agreement[14],
                '@ri_iban': agreement[15],
                '@bank_account_detail_ri': agreement[16],
                '@ri_shortname': agreement[17]
            }

            # Замена текста в шаблоне
            self.replace_text_in_document(doc, replacements)
            self.replace_in_tables(doc.tables, replacements)
            self.formatting_text(doc)

            # Сохранение документа
            doc_io = BytesIO()
            doc.save(doc_io)
            doc_io.seek(0)

            # Формируем название файла
            file_name = f"{agreement[1]}_Договір.docx"

            # Отправка документа клиенту
            return await send_file(doc_io, as_attachment=True, attachment_filename=file_name)

        @self.app.route('/llc_acts/<int:agreement_id>/generate_llc_appendix', methods=['GET'])
        @basic_auth_required()
        async def generate_llc_appendix(agreement_id):
            # Получаем данные по договору
            agreement_query = """
            SELECT la.id, la.agreement_name, la.agreement_date, 
                   lc.name AS llc_name, lc.edrpou AS llc_edrpou, 
                   lc.address AS llc_address, lc.iban AS llc_iban, lc.bank_account_detail AS bank_account_detail_llc, 
                   lc.name_short AS llc_shortname, ri.id AS engineer_id, lc.id AS llc_id, lc.in_persona, ri.name, ri.inn,
                   ri.pidstava, ri.address AS ri_address, ri.iban, ri.bank_account_detail, ri.name_short, lc.inn
            FROM credentials.llc_agreements AS la
            JOIN credentials.llc_credentials AS lc ON la.llc_id = lc.id
            JOIN credentials.ri_credentials AS ri ON la.ri_id = ri.id
            WHERE la.id = %s;
            """
            agreement_data = await self.local_db.execute_query(agreement_query, (agreement_id,))
            print(agreement_data)

            if not agreement_data:
                return "Договор не найден", 404

            agreement = agreement_data[0]
            llc_edrpou = agreement[4]

            # Выбираем запросы в зависимости от ЕДРПОУ
            if llc_edrpou == 38736443:
                model_query = """
                SELECT sw.model, COUNT(DISTINCT sw.ip)
                FROM dbsyphon.switches_report sw
                JOIN credentials.engineer_cantons ct ON sw.canton = ct.canton
                WHERE sw.switch_rank IN (1,2) AND ct.engineer_id = %s
                GROUP BY 1;
                """
                ip_pool_query = """
                SELECT CONCAT(SUBSTRING_INDEX(sr.ip, '.', 3), '.0/24') AS ip_pool
                FROM dbsyphon.switches_report sr
                JOIN credentials.engineer_cantons ct ON sr.canton = ct.canton
                WHERE sr.switch_rank IN (1,2) AND ct.engineer_id = %s
                GROUP BY ip_pool;
                """
                query_param = agreement[9]  # engineer_id
            else:
                model_query = """
                SELECT sw.model, COUNT(DISTINCT sw.ip)
                FROM dbsyphon.switches_report sw
                JOIN credentials.llc_cantons ct ON sw.canton = ct.canton
                WHERE sw.switch_rank IN (3,4) AND ct.llc_id = %s
                GROUP BY 1;
                """
                ip_pool_query = """
                SELECT CONCAT(SUBSTRING_INDEX(sr.ip, '.', 3), '.0/24') AS ip_pool
                FROM dbsyphon.switches_report sr
                JOIN credentials.llc_cantons ct ON sr.canton = ct.canton
                WHERE sr.switch_rank IN (3,4) AND ct.llc_id = %s
                GROUP BY ip_pool;
                """
                query_param = agreement[10]  # llc_id

            # Выполняем запросы
            data_table1 = await self.local_db.execute_query(model_query, (query_param,))
            data_table2 = await self.local_db.execute_query(ip_pool_query, (query_param,))

            # Загружаем шаблон документа
            template_path = 'static/docs/llc_appendix.docx'
            doc = Document(template_path)

            # Форматируем дату договора
            agreement_date_str, month_ukr_name, year, _ = self.format_date(agreement[2])

            # Заполняем заменяемые поля
            replacements = {
                '@agr_num': agreement[0],
                '@agr_name': agreement[1],
                '@agr_date': agreement_date_str,
                '@llc_name': agreement[3],
                '@persona': agreement[11],
                '@ri_name': agreement[12],
                '@ri_inn': agreement[13],
                '@pidstava': agreement[14],
                '@ri_address': agreement[15],
                '@ri_iban': agreement[16],
                '@llc_inn': agreement[19],
                '@bank_account_detail_ri': agreement[17],
                '@ri_shortname': agreement[18],
                '@llc_edrpou': llc_edrpou,
                '@llc_address': agreement[5],
                '@llc_iban': agreement[6],
                '@bank_account_detail_llc': agreement[7],
                '@llc_shortname': agreement[8]
            }

            # Выполняем замену в шаблоне
            self.replace_text_in_document(doc, replacements)
            self.replace_in_tables(doc.tables, replacements)
            self.formatting_text(doc)

            # Создаём таблицы
            table1 = self.create_table(doc, data_table1, ['Модель обладнання', 'Кількість'])
            table2 = self.create_table(doc, data_table2, ['Діапазон ІР адрес'])

            # Вставляем таблицы в документ
            self.replace_table_in_document(doc, '@table1', table1)
            self.replace_table_in_document(doc, '@table2', table2)

            # Сохраняем документ в память
            doc_io = BytesIO()
            doc.save(doc_io)
            doc_io.seek(0)

            # Формируем название файла
            file_name = f"{agreement[1]}_Додаток.docx"

            # Отправляем файл клиенту
            return await send_file(doc_io, as_attachment=True, attachment_filename=file_name)

        @self.app.route('/llc_acts/<int:agreement_id>/delete/<int:act_id>', methods=['POST'])
        @basic_auth_required()
        async def delete_act(agreement_id, act_id):
            # Удаление акта из базы данных
            delete_query = """
            UPDATE credentials.llc_acts SET act_state=0 
            where id = %s AND agreement = %s
            """
            await self.local_db.execute_query(delete_query, (act_id, agreement_id))

            # Перенаправление обратно на страницу актов
            return redirect(url_for('llc_acts', agreement_id=agreement_id))

        @self.app.route('/llc_acts/<int:agreement_id>', methods=['GET', 'POST'])
        @basic_auth_required()
        async def llc_acts(agreement_id):
            if request.method == 'POST':
                # Получение данных из формы
                act_date = (await request.form)['act_date']
                act_sum = (await request.form)['act_sum']

                # Вставка нового акта в базу данных
                insert_query = """
                INSERT INTO credentials.llc_acts (agreement, act_date, act_sum, act_state)
                VALUES (%s, %s, %s, 1);
                """
                await self.local_db.execute_query(insert_query, (agreement_id, act_date, act_sum))

                # Перезагрузка страницы после добавления акта
                return redirect(url_for('llc_acts', agreement_id=agreement_id))

            # Запрос данных по договору и связанных таблиц
            agreement_query = """
            SELECT la.agreement_name, lc.name AS organization_name, lc.edrpou, ri.name AS engineer_name
            FROM credentials.llc_agreements AS la
            JOIN credentials.llc_credentials AS lc ON la.llc_id = lc.id
            JOIN credentials.ri_credentials AS ri ON la.ri_id = ri.id
            WHERE la.id = %s;
            """
            agreement_data = await self.local_db.execute_query(agreement_query, (agreement_id,))

            if not agreement_data:
                return "Договор не найден", 404

            agreement = {
                'agreement_name': agreement_data[0][0],
                'organization_name': agreement_data[0][1],
                'edrpou': agreement_data[0][2],
                'engineer_name': agreement_data[0][3]
            }

            # Запрос данных об актах
            acts_query = """
            SELECT id, act_date, act_sum, act_state
            FROM credentials.llc_acts
            WHERE agreement = %s AND act_state = 1;
            """
            acts_data = await self.local_db.execute_query(acts_query, (agreement_id,))

            # Проверка данных в таблице llc_acts_data
            if acts_data:
                act_ids = [act[0] for act in acts_data]
                acts_with_data_query = """
                SELECT DISTINCT act_id FROM credentials.llc_acts_data WHERE act_id IN (%s);
                """
                # Генерация списка идентификаторов для подстановки в запрос
                in_clause = ','.join(['%s'] * len(act_ids))
                acts_with_data_query = acts_with_data_query % in_clause

                acts_with_data = await self.local_db.execute_query(acts_with_data_query, tuple(act_ids))
                acts_with_data_ids = {row[0] for row in acts_with_data}
            else:
                acts_with_data_ids = set()

            # Преобразование данных о актах в список словарей
            acts = [
                {
                    'id': act[0],
                    'act_date': act[1],
                    'act_sum': act[2],
                    'act_state': act[3],
                    'has_data': act[0] in acts_with_data_ids
                }
                for act in acts_data
            ]

            return await render_template('llc_acts.html', agreement=agreement, acts=acts, agreement_id=agreement_id)

        @self.app.route('/kdn-new/acts/<int:agreement_id>', methods=['GET', 'POST'])
        @basic_auth_required()
        async def kdn_new_acts(agreement_id):
            if request.method == 'POST':
                act_date = (await request.form)['act_date']
                act_sum = (await request.form)['act_sum']

                insert_query = """
                INSERT INTO credentials.llc_acts_new (agreement, act_date, act_sum, act_state)
                VALUES (%s, %s, %s, 1);
                """
                await self.local_db.execute_query(insert_query, (agreement_id, act_date, act_sum))

                return redirect(url_for('kdn_new_acts', agreement_id=agreement_id))

            agreement_query = """
            SELECT la.agreement_name, lc.name AS organization_name, lc.edrpou, ri.name AS engineer_name
            FROM credentials.llc_agreements_new AS la
            JOIN credentials.llc_credentials AS lc ON la.llc_id = lc.id
            JOIN credentials.ri_credentials AS ri ON la.ri_id = ri.id
            WHERE la.id = %s;
            """
            agreement_data = await self.local_db.execute_query(agreement_query, (agreement_id,))

            if not agreement_data:
                return "Договор не найден", 404

            agreement = {
                'agreement_name': agreement_data[0][0],
                'organization_name': agreement_data[0][1],
                'edrpou': agreement_data[0][2],
                'engineer_name': agreement_data[0][3]
            }

            acts_query = """
            SELECT id, act_date, act_sum, act_state
            FROM credentials.llc_acts_new
            WHERE agreement = %s AND act_state = 1;
            """
            acts_data = await self.local_db.execute_query(acts_query, (agreement_id,))

            if acts_data:
                act_ids = [act[0] for act in acts_data]
                placeholders = build_placeholders(act_ids)
                acts_with_data_query = f"""
                SELECT DISTINCT act_id
                FROM credentials.llc_acts_data_new
                WHERE act_id IN ({placeholders});
                """
                acts_with_data = await self.local_db.execute_query(acts_with_data_query, tuple(act_ids))
                acts_with_data_ids = {row[0] for row in acts_with_data}
            else:
                acts_with_data_ids = set()

            acts = [
                {
                    'id': act[0],
                    'act_date': act[1],
                    'act_sum': act[2],
                    'act_state': act[3],
                    'has_data': act[0] in acts_with_data_ids,
                }
                for act in acts_data
            ]

            return await render_template('kdn_new_acts.html', agreement=agreement, acts=acts, agreement_id=agreement_id)

        async def get_kdn_new_act_document_data(act_id):
            query = """
            SELECT act.act_date,
                   act.act_sum,
                   act.id,
                   agr.agreement_name,
                   agr.agreement_date,
                   llc.name,
                   llc.edrpou,
                   ri.name,
                   ri.iban,
                   ri.bank_account_detail,
                   ri.address,
                   ri.phone,
                   ri.inn,
                   ri.name_short,
                   llc.in_persona,
                   ri.pidstava,
                   llc.address,
                   llc.iban,
                   llc.bank_account_detail,
                   llc.inn,
                   llc.name_short
            FROM credentials.llc_acts_new act
            JOIN credentials.llc_agreements_new agr ON act.agreement = agr.id
            JOIN credentials.llc_credentials llc ON agr.llc_id = llc.id
            JOIN credentials.ri_credentials ri ON agr.ri_id = ri.id
            WHERE act.id = %s
            """
            result = await self.local_db.execute_query(query, (act_id,))
            data = result[0] if result else None

            query_acts_data = """
            SELECT sw_rank, model_list, count_devices, ip_list, worktime_float
            FROM credentials.llc_acts_data_new
            WHERE act_id = %s
            """
            acts_data_result = await self.local_db.execute_query(query_acts_data, (act_id,))
            return data, acts_data_result

        @self.app.route('/kdn-new/acts/<int:agreement_id>/generate_data/<int:act_id>', methods=['POST'])
        @basic_auth_required()
        async def generate_kdn_new_act_data(agreement_id, act_id):
            act_query = """
            SELECT act_date, act_sum, agreement
            FROM credentials.llc_acts_new
            WHERE id = %s;
            """
            act_data = await self.local_db.execute_query(act_query, (act_id,))

            if not act_data:
                return "Акт не найден", 404

            act_date, act_sum, agreement = act_data[0]

            agreement_query = """
            SELECT llc_id, ri_id
            FROM credentials.llc_agreements_new
            WHERE id = %s;
            """
            agreement_data = await self.local_db.execute_query(agreement_query, (agreement,))

            if not agreement_data:
                return "Договор не найден", 404

            llc_id, ri_id = agreement_data[0]

            edrpou_query = """
            SELECT edrpou
            FROM credentials.llc_credentials
            WHERE id = %s;
            """
            edrpou_data = await self.local_db.execute_query(edrpou_query, (llc_id,))

            if not edrpou_data:
                return "Организация не найдена", 404

            if edrpou_data[0][0] == 38736443:
                await handle_kdn_logic(
                    act_date,
                    act_sum,
                    agreement,
                    ri_id,
                    act_id,
                    "credentials.llc_acts_data_new",
                    "credentials.engineer_cantons_new"
                )
            else:
                await handle_llc_logic(act_sum, llc_id, act_id, "credentials.llc_acts_data_new")

            return redirect(url_for('kdn_new_acts', agreement_id=agreement_id))

        @self.app.route('/kdn-new/acts/<int:act_id>/generate_report_llc', methods=['POST'])
        @basic_auth_required()
        async def generate_kdn_new_report_llc(act_id):
            data, acts_data_result = await get_kdn_new_act_document_data(act_id)
            if not data:
                return "Акт не найден", 404

            replacements = {
                "@act_name": f"R{act_id}_{data[0].strftime('%m/%y')}_{data[3]}",
                "@act_date": f"«{calendar.monthrange(int(data[0].strftime('%Y')), int(data[0].month))[-1]}» {self.format_date(data[0])[1]} {self.format_date(data[0])[2]} року",
                "@ri_name": data[7],
                "@ri_iban": data[8],
                "@bank_account_detail_ri": data[9],
                "@ri_address": data[10],
                "@ri_phone": data[11],
                "@ri_inn": data[12],
                "@llc_name": data[5],
                "@llc_edrpou": data[6],
                "@agr_name": data[3],
                "@agr_date": self.format_date(data[4])[0],
                "@act_sum": data[1],
                "@actsumwords": self.convert_to_currency_words(data[1]),
                "@ri_shortname": data[13],
                "@llc_in_persona": data[14],
                "@ri_pidstava": data[15],
                "@current_month": self.format_date(data[0])[1],
                "@current_year": self.format_date(data[0])[2],
                "@last_day_of_the_month": calendar.monthrange(int(data[0].strftime('%Y')), int(data[0].month))[-1],
                "@llc_address": data[16],
                "@llc_iban": data[17],
                "@bank_account_detail_llc": data[18],
                "@llc_inn": data[19],
                "@llc_shortname": data[20]
            }

            for row in acts_data_result:
                if row[0] == 1:
                    replacements["@rank1_models"] = row[1]
                    replacements["@rank1_count"] = row[2]
                    replacements["@rank1_ips"] = row[3]
                elif row[0] == 2:
                    replacements["@rank2_models"] = row[1]
                    replacements["@rank2_count"] = row[2]
                    replacements["@rank2_ips"] = row[3]
                elif row[0] == 0:
                    replacements["@time_models"] = row[1]
                    replacements["@time_count"] = round(row[4], 2)
                    replacements["@time_ips"] = row[3]

            document = Document('static/docs/kdn_report.docx')
            self.replace_text_in_document(document, replacements)
            self.replace_in_tables(document.tables, replacements)
            self.formatting_text(document)
            clear_document_highlights(document)

            output = BytesIO()
            document.save(output)
            pdf_bytes = await convert_office_bytes_to_pdf(output.getvalue(), ".docx")
            docx_name = f'{data[3]}_Звіт_{self.format_date(data[0])[1]}_{self.format_date(data[0])[2]}'
            return await send_file(BytesIO(pdf_bytes), as_attachment=True, attachment_filename=f"{docx_name}.pdf",
                                   mimetype="application/pdf")

        @self.app.route('/kdn-new/acts/<int:act_id>/generate_act', methods=['POST'])
        @basic_auth_required()
        async def generate_kdn_new_act(act_id):
            data, acts_data_result = await get_kdn_new_act_document_data(act_id)
            if not data:
                return "Акт не найден", 404

            replacements = {
                "@act_name": f"A{act_id}_{data[0].strftime('%m/%y')}_{data[3]}",
                "@act_date": f"«{calendar.monthrange(int(data[0].strftime('%Y')), int(data[0].month))[-1]}» {self.format_date(data[0])[1]} {self.format_date(data[0])[2]} року",
                "@ri_name": data[7],
                "@ri_iban": data[8],
                "@bank_account_detail_ri": data[9],
                "@ri_address": data[10],
                "@ri_phone": data[11],
                "@ri_inn": data[12],
                "@llc_name": data[5],
                "@llc_edrpou": data[6],
                "@agr_name": data[3],
                "@agr_date": self.format_date(data[4])[0],
                "@act_sum": data[1],
                "@actsumwords": self.convert_to_currency_words(data[1]),
                "@ri_shortname": data[13],
                "@llc_in_persona": data[14],
                "@ri_pidstava": data[15],
                "@current_month": self.format_date(data[0])[1],
                "@current_year": self.format_date(data[0])[2],
                "@last_day_of_the_month": calendar.monthrange(int(data[0].strftime('%Y')), int(data[0].month))[-1],
                "@llc_address": data[16],
                "@llc_iban": data[17],
                "@bank_account_detail_llc": data[18],
                "@llc_inn": data[19],
                "@llc_shortname": data[20]
            }
            sum_rank1 = 0
            sum_rank2 = 0

            for row in acts_data_result:
                if row[0] == 1:
                    replacements["@rank1_count"] = row[2]
                    replacements["@rank1_sum"] = float(row[2] * 1000)
                    sum_rank1 = float(row[2] * 1000)
                elif row[0] == 2:
                    replacements["@rank2_count"] = row[2]
                    replacements["@rank2_sum"] = float(row[2] * 1000)
                    sum_rank2 = float(row[2] * 1000)
                elif row[0] == 0:
                    replacements["@time_count"] = round(row[4], 2)
                    replacements["@time_sum"] = round(float(data[1]) - sum_rank2 - sum_rank1, 2)

            document = Document('static/docs/kdn_act.docx')
            self.replace_text_in_document(document, replacements)
            self.replace_in_tables(document.tables, replacements)
            self.formatting_text(document)
            clear_document_highlights(document)

            output = BytesIO()
            document.save(output)
            pdf_bytes = await convert_office_bytes_to_pdf(output.getvalue(), ".docx")
            docx_name = f'{data[3]}_Акт_{self.format_date(data[0])[1]}_{self.format_date(data[0])[2]}'
            return await send_file(BytesIO(pdf_bytes), as_attachment=True, attachment_filename=f"{docx_name}.pdf",
                                   mimetype="application/pdf")

        @self.app.route('/kdn-new/acts/<int:act_id>/generate_bill', methods=['POST'])
        @basic_auth_required()
        async def generate_kdn_new_bill(act_id):
            data, acts_data_result = await get_kdn_new_act_document_data(act_id)
            if not data:
                return "Акт не найден", 404

            replacements = {
                "@bill_name": f"B{act_id}_{data[0].strftime('%m/%y')}_{data[3]}",
                "@bill_date": f"{calendar.monthrange(int(data[0].strftime('%Y')), int(data[0].month))[-1]} {self.format_date(data[0])[1]} {self.format_date(data[0])[2]} року",
                "@ri_name": data[7],
                "@ri_iban": data[8],
                "@bank_account_detail_ri": data[9],
                "@ri_address": data[10],
                "@ri_phone": data[11],
                "@ri_inn": data[12],
                "@llc_name": data[5],
                "@llc_edrpou": data[6],
                "@agr_name": data[3],
                "@agr_date": self.format_date(data[4])[0],
                "@bill_sum": data[1],
                "@handwritebill_sum": self.convert_to_currency_words(data[1]),
                "@ri_shortname": data[13]
            }
            sum_rank1 = 0
            sum_rank2 = 0

            for row in acts_data_result:
                if row[0] == 1:
                    replacements["@rank1_count"] = row[2]
                    replacements["@rank1_sum"] = float(row[2] * 1000)
                    sum_rank1 = float(row[2] * 1000)
                elif row[0] == 2:
                    replacements["@rank2_count"] = row[2]
                    replacements["@rank2_sum"] = float(row[2] * 1000)
                    sum_rank2 = float(row[2] * 1000)
                elif row[0] == 0:
                    replacements["@time_count"] = round(row[4], 2)
                    replacements["@time_sum"] = round(float(data[1]) - sum_rank2 - sum_rank1, 2)

            workbook = openpyxl.load_workbook('static/docs/kdn_bill.xlsx')
            sheet = workbook.active

            for row in sheet.iter_rows():
                for cell in row:
                    if isinstance(cell.value, str):
                        for key, replacement in replacements.items():
                            if key in cell.value:
                                cell.value = cell.value.replace(key, str(replacement))
            clear_workbook_highlights(workbook)

            output = BytesIO()
            workbook.save(output)
            pdf_bytes = await convert_office_bytes_to_pdf(output.getvalue(), ".xlsx")
            xlxs_name = f'{data[3]}_Рахунок_{self.format_date(data[0])[1]}_{self.format_date(data[0])[2]}'
            return await send_file(BytesIO(pdf_bytes), as_attachment=True, attachment_filename=f"{xlxs_name}.pdf",
                                   mimetype="application/pdf")

        @self.app.route('/kdn-new/acts/<int:agreement_id>/delete/<int:act_id>', methods=['POST'])
        @basic_auth_required()
        async def delete_kdn_new_act(agreement_id, act_id):
            delete_query = """
            UPDATE credentials.llc_acts_new
            SET act_state = 0
            WHERE id = %s AND agreement = %s;
            """
            await self.local_db.execute_query(delete_query, (act_id, agreement_id))
            return redirect(url_for('kdn_new_acts', agreement_id=agreement_id))


        @self.app.route('/llc_agreements', methods=['GET'])
        @basic_auth_required()
        async def llc_agreements():
            # SQL-запрос для получения информации о всех договорах
            query = """
            SELECT la.id, la.agreement_name, la.agreement_date, llc.name AS llc_name, llc.canton AS canton, ri.name AS engineer_name
            FROM credentials.llc_agreements la
            JOIN credentials.llc_credentials llc ON la.llc_id = llc.id
            JOIN credentials.ri_credentials ri ON la.ri_id = ri.id
            WHERE la.agreement_state = 1;
            """

            # Выполняем запрос и получаем данные по договорам
            agreements_data = await self.local_db.execute_query(query)

            all_years = set()
            agreements_list = []

            # Преобразуем каждый кортеж в словарь
            for agreement in agreements_data:
                agreement_dict = {
                    'id': agreement[0],
                    'agreement_name': agreement[1],
                    'agreement_date': agreement[2],
                    'llc_name': agreement[3],
                    'canton': agreement[4],
                    'engineer_name': agreement[5],
                    'acts_by_year': {}
                }

                # Для каждого договора получаем данные по актам
                act_query = """
                SELECT YEAR(act_date) AS act_year, MONTH(act_date) AS act_month
                FROM credentials.llc_acts
                WHERE act_state = 1 AND agreement = %s
                ORDER BY act_year, act_month;
                """
                act_data = await self.local_db.execute_query(act_query, (agreement[0],))

                # Сортируем акты по годам
                for row in act_data:
                    year, month = row
                    if year not in agreement_dict['acts_by_year']:
                        agreement_dict['acts_by_year'][year] = []
                    agreement_dict['acts_by_year'][year].append(month)
                    all_years.add(year)

                agreements_list.append(agreement_dict)

            # Передаем все данные в шаблон, включая все уникальные годы
            return await render_template('llc_agreements.html', agreements=agreements_list, all_years=sorted(all_years))

        @self.app.route('/kdn-new', methods=['GET'])
        @basic_auth_required()
        async def kdn_new():
            query = """
            SELECT la.id, la.agreement_name, la.agreement_date, llc.name AS llc_name, llc.canton AS canton, ri.name AS engineer_name
            FROM credentials.llc_agreements_new la
            JOIN credentials.llc_credentials llc ON la.llc_id = llc.id
            JOIN credentials.ri_credentials ri ON la.ri_id = ri.id
            WHERE la.agreement_state = 1 AND llc.edrpou = %s;
            """

            agreements_data = await self.local_db.execute_query(query, (38736443,))

            all_years = set()
            agreements_list = []

            for agreement in agreements_data:
                agreement_dict = {
                    'id': agreement[0],
                    'agreement_name': agreement[1],
                    'agreement_date': agreement[2],
                    'llc_name': agreement[3],
                    'canton': agreement[4],
                    'engineer_name': agreement[5],
                    'acts_by_year': {}
                }

                act_query = """
                SELECT YEAR(act_date) AS act_year, MONTH(act_date) AS act_month
                FROM credentials.llc_acts_new
                WHERE act_state = 1 AND agreement = %s
                ORDER BY act_year, act_month;
                """
                act_data = await self.local_db.execute_query(act_query, (agreement[0],))

                for row in act_data:
                    year, month = row
                    if year not in agreement_dict['acts_by_year']:
                        agreement_dict['acts_by_year'][year] = []
                    agreement_dict['acts_by_year'][year].append(month)
                    all_years.add(year)

                agreements_list.append(agreement_dict)

            return await render_template('kdn_new.html', agreements=agreements_list, all_years=sorted(all_years))

        @self.app.route('/correct_agreement/<int:id>', methods=['POST'])
        @basic_auth_required()
        async def correct_agreement(id):
            try:
                # Получаем данные из таблицы protocols_missing_agreements по id
                missing_agreement_query = """
                    SELECT fop_inn, ri_inn, date_of_protocol, fop_change 
                    FROM credentials.protocols_missing_agreements 
                    WHERE id = %s;
                """
                missing_agreement = await self.local_db.execute_query(missing_agreement_query, (id,))

                if not missing_agreement:
                    await flash("Недостающий договор не найден.", "error")
                    return redirect(url_for('missing_agreements'))

                fop_inn, ri_inn, date_of_protocol, fop_change = missing_agreement[0]

                # Проверяем наличие договора в таблице agreements
                agreement_query = """
                    SELECT agreements.id 
                    FROM credentials.agreements AS agreements
                    JOIN credentials.fop_credentials AS fop ON agreements.master_id = fop.id
                    JOIN credentials.ri_credentials AS ri ON agreements.ri_id = ri.id
                    WHERE fop.inn = %s AND ri.inn = %s;
                """
                agreement_result = await self.local_db.execute_query(agreement_query, (fop_inn, ri_inn))
                agreement_id = agreement_result[0][0] if agreement_result else None

                if agreement_id:
                    # Вставляем данные в таблицу protocols_test
                    proto_sum_caps = self.convert_to_currency_words(fop_change)

                    insert_protocol_query = """
                        INSERT INTO credentials.protocols (agreement, proto_date, proto_sum, proto_sum_caps)
                        VALUES (%s, %s, %s, %s);
                    """
                    await self.local_db.execute_query(insert_protocol_query,
                                                      (agreement_id, date_of_protocol, fop_change, proto_sum_caps))

                    # Обновляем состояние в таблице protocols_missing_agreements
                    update_missing_agreement_query = """
                        UPDATE credentials.protocols_missing_agreements 
                        SET agreement_state = 1 
                        WHERE id = %s;
                    """
                    await self.local_db.execute_query(update_missing_agreement_query, (id,))

                    await flash("Договор успешно исправлен и протокол добавлен.", "success")
                    return redirect(url_for('missing_agreements'))
                else:
                    await flash("Договор не найден.", "error")
                    return redirect(url_for('missing_agreements'))

            except Exception as e:
                await flash(f"Ошибка при исправлении договора: {e}", "error")
                return redirect(url_for('missing_agreements'))

        @self.app.route('/missing_agreements', methods=['GET'])
        @basic_auth_required()
        async def missing_agreements():
            try:
                # Получаем записи с agreement_state = 0
                missing_agreements_query = """
                    SELECT id, clientId, description, fop_inn, fop_name, fop_in, fop_change, fop_expense, fop_out, 
                           type_agr, ri_inn, ri_name, date_of_protocol
                    FROM credentials.protocols_missing_agreements
                    WHERE agreement_state = 0;
                """
                agreements = await self.local_db.execute_query(missing_agreements_query)

                # Возвращаем HTML-страницу с данными
                return await render_template('missing_agreements.html', agreements=agreements)

            except Exception as e:
                await flash(f"Ошибка при загрузке недостающих договоров: {e}", "error")

        @self.app.route('/agreement_insertion', methods=['GET'])
        @basic_auth_required()
        async def agreement_insertion():
            # Загружаем всех инженеров для отображения в выпадающем списке
            engineers_query = "SELECT id, name FROM credentials.ri_credentials;"
            engineers = await self.local_db.execute_query(engineers_query)

            # Отправляем список инженеров на страницу
            return await render_template("agreement_insertion.html", engineers=engineers)

        @self.app.route('/search_masters', methods=['GET'])
        @basic_auth_required()
        async def search_masters():
            query = request.args.get("query", "")
            search_query = """
                SELECT id, name FROM credentials.fop_credentials 
                WHERE name LIKE %s LIMIT 10;
            """
            results = await self.local_db.execute_query(search_query, (f"%{query}%",))
            return jsonify([{"id": result[0], "name": result[1]} for result in results])

        @self.app.route('/submit_agreement', methods=['POST'])
        @basic_auth_required()
        async def submit_agreement():
            try:
                # Используем await перед request.form для корректного извлечения данных формы
                form_data = await request.form

                agreement_name = form_data.get('agreement_name')
                master_id = form_data.get('master_id')
                engineer_id = form_data.get('engineer')
                agreement_date = form_data.get('agreement_date')

                # Вставка данных в таблицу agreements
                insert_query = """
                    INSERT INTO credentials.agreements (agreement_name, master_id, ri_id, agreement_date)
                    VALUES (%s, %s, %s, %s);
                """
                await self.local_db.execute_query(insert_query,
                                                  (agreement_name, master_id, engineer_id, agreement_date))

                await flash("Новый договор успешно добавлен!", "success")
            except Exception as e:
                await flash(f"Ошибка при добавлении договора: {e}", "error")

            # Перенаправление на страницу добавления договора
            return redirect(url_for('agreement_insertion'))

        @self.app.route('/get_master_inn/<int:master_id>', methods=['GET'])
        @basic_auth_required()
        async def get_master_inn(master_id):
            """
            Получает ИНН мастера по его ID.
            """
            query = "SELECT inn FROM credentials.fop_credentials WHERE id = %s;"
            result = await self.local_db.execute_query(query, (master_id,))

            if result and result[0]:
                return jsonify({"inn": result[0][0]}), 200
            else:
                return jsonify({"error": "Мастер не найден"}), 404

        @self.app.route('/generate_protocols', methods=['POST'])
        @basic_auth_required()
        async def generate_protocols():
            try:
                # Получаем месяц и год из формы
                protocol_month = int((await request.form).get('protocol_month'))
                protocol_year = int((await request.form).get('protocol_year'))

                # Запрос только для записей с указанным месяцем и годом
                soft_estimates_query = """
                    SELECT id, clientId, description, fop_inn, fop_name, fop_in, fop_change, 
                           fop_expense, fop_out, type_agr, ri_inn, ri_name, date_of_protocol
                    FROM credentials.soft_estimates
                    WHERE MONTH(date_of_protocol) = %s AND YEAR(date_of_protocol) = %s;
                """

                # Выполняем запрос
                soft_estimates = await self.local_db.execute_query(soft_estimates_query,
                                                                   (protocol_month, protocol_year))
                print("Количество записей для генерации протоколов:", len(soft_estimates))

                if not soft_estimates:
                    await flash("Записей для указанного месяца и года не найдено.", "info")
                    return redirect(url_for('estimates_upload'))

                for record in soft_estimates:
                    (id, clientId, description, fop_inn, fop_name, fop_in, fop_change,
                     fop_expense, fop_out, type_agr, ri_inn, ri_name, date_of_protocol) = record

                    # Поиск соответствующего договора
                    agreement_query = """
                        SELECT agreements.id 
                        FROM credentials.agreements AS agreements
                        JOIN credentials.fop_credentials AS fop ON agreements.master_id = fop.id
                        JOIN credentials.ri_credentials AS ri ON agreements.ri_id = ri.id
                        WHERE fop.inn = %s AND ri.inn = %s;
                    """
                    agreement_result = await self.local_db.execute_query(agreement_query, (fop_inn, ri_inn))
                    agreement = agreement_result[0][0] if agreement_result else None
                    print("Договор найден:", agreement)

                    if agreement:
                        # Договор найден, вставляем данные в таблицу protocols
                        proto_sum_caps = self.convert_to_currency_words(fop_change)

                        insert_protocol_query = """
                            INSERT INTO credentials.protocols (agreement, proto_date, proto_sum, proto_sum_caps)
                            VALUES (%s, %s, %s, %s);
                        """
                        await self.local_db.execute_query(insert_protocol_query,
                                                          (agreement, date_of_protocol, fop_change, proto_sum_caps))
                        print(f"Протокол добавлен для договора {agreement} на сумму {fop_change}")
                    else:
                        # Договор не найден, копируем все данные в protocols_missing_agreements
                        insert_missing_agreement_query = """
                            INSERT INTO credentials.protocols_missing_agreements 
                            (clientId, description, fop_inn, fop_name, fop_in, fop_change, fop_expense, fop_out, 
                             type_agr, ri_inn, ri_name, date_of_protocol, agreement_state)
                            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s);
                        """
                        await self.local_db.execute_query(insert_missing_agreement_query,
                                                          (clientId, description, fop_inn, fop_name, fop_in, fop_change,
                                                           fop_expense, fop_out, type_agr, ri_inn, ri_name,
                                                           date_of_protocol, False))
                        print(
                            f"Протокол не добавлен, данные сохранены в таблицу protocols_missing_agreements для клиента {clientId}")

                await flash("Протоколы успешно сгенерированы и сохранены.", "success")

            except Exception as e:
                print(f"Ошибка при генерации протоколов: {e}")
                await flash(f"Ошибка при генерации протоколов: {e}", "error")

            # Перенаправляем пользователя на нужную страницу
            return redirect(url_for('estimates_upload'))

        @self.app.route('/estimates_upload', methods=['GET', 'POST'])
        @basic_auth_required()
        async def estimates_upload():
            # Обработка POST запроса для загрузки данных
            if request.method == 'POST':
                # Получение данных из формы с использованием await
                date_str = (await request.form).get('date')
                file = (await request.files).get('file')

                if not date_str or not file:
                    await flash("Пожалуйста, укажите дату и выберите файл.")
                    return redirect(url_for('estimates_upload'))

                # Преобразуем дату в нужный формат
                try:
                    date_of_protocol = datetime.strptime(date_str, '%Y-%m-%d').date()
                except ValueError:
                    await flash("Некорректный формат даты.")
                    return redirect(url_for('estimates_upload'))

                # Чтение данных из XLSX
                df = pd.read_excel(file)
                insert_query = """
                    INSERT INTO credentials.soft_estimates (
                        clientId, description, fop_inn, fop_name, fop_in,
                        fop_change, fop_expense, fop_out, type_agr, 
                        ri_inn, ri_name, date_of_protocol
                    ) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                """

                # Обработка данных и вставка в базу
                records = [
                    (
                        int(row['ClientId']), str(row['Description']), int(row['OKPO']), str(row['Name']),
                        float(row['In']), float(row['Charge']), float(row['Expense']),
                        float(row['Out']), str(row['Type']), int(row['ContractOKPO']),
                        str(row['ContractName']), date_of_protocol
                    )
                    for _, row in df.iterrows()
                ]

                try:
                    # Выполнение вставки
                    await self.local_db.ensure_connection()
                    for record in records:
                        await self.local_db.execute_query(insert_query, record)
                    await flash("Данные успешно загружены.")
                except Exception as e:
                    await flash(f"Ошибка при загрузке данных: {e}")
                    print(f"Ошибка при загрузке данных: {e}")

                return redirect(url_for('estimates_upload'))

            # Обработка GET запроса для отображения формы и уникальных дат
            select_dates_query = "SELECT DISTINCT date_of_protocol FROM credentials.soft_estimates"
            try:
                await self.local_db.ensure_connection()
                dates = await self.local_db.execute_query(select_dates_query)
            except Exception as e:
                print(f"Ошибка при получении дат: {e}")
                dates = []

            return await render_template('estimates_upload.html', dates=dates)


        @self.app.route('/protocols/<int:agreement_id>/generate_docx/<int:protocol_id>', methods=['GET'])
        @basic_auth_required()
        async def generate_docx(agreement_id,protocol_id):
            # Получаем информацию по договору и протоколу
            agreement_query = """
            SELECT a.agreement_name, a.agreement_date, f.name AS fop_name, f.inn AS inn_fop, f.pidstava AS pidstava_fop, 
                   f.address AS fop_address, f.iban AS fop_iban, f.bank_account_detail AS bank_account_detail_fop, f.name_short AS fop_name_short,
                   r.name AS ri_name, r.inn AS inn_ri, r.pidstava AS pidstava_ri, r.address AS ri_address, r.iban AS ri_iban, r.bank_account_detail AS bank_account_detail_ri, r.name_short AS ri_name_short
            FROM credentials.agreements AS a
            JOIN credentials.fop_credentials AS f ON a.master_id = f.id
            JOIN credentials.ri_credentials AS r ON a.ri_id = r.id
            WHERE a.id = %s;
            """
            agreement_data = await self.local_db.execute_query(agreement_query, (agreement_id,))
            agreement = agreement_data[0]  # Мы получаем первый элемент, чтобы передать данные как строку, а не кортеж.

            protocol_query = """
                        SELECT proto_date, proto_sum, proto_sum_caps
                        FROM credentials.protocols
                        WHERE agreement = %s and id = %s;
                        """
            protocol_data = await self.local_db.execute_query(protocol_query, (agreement_id, protocol_id))
            protocol = protocol_data[0]  # Тоже получаем первый элемент из протоколов

            if not agreement_data or not protocol_data:
                return "Договор или протокол не найден", 404

            # Преобразуем данные и форматируем дату
            agreement_date_str, month_ukr_name, year, _ = self.format_date(agreement[1])
            proto_date_str, proto_month_ukr_name, proto_year, last_day_of_the_month  = self.format_date(protocol[0])
            template_month = calendar.monthrange(int(protocol[0].strftime("%Y")), int(protocol[0].month))
            last_day_of_the_month = str(template_month[1])

            # Загрузка шаблона
            template_path = 'static/docs/M-RI_protocol.docx'
            doc = Document(template_path)

            # Замена маркеров
            replacements = {
                '@agr_num': agreement[0],
                '@agr_date': agreement_date_str,
                '@proto_date': proto_date_str,
                '@fop_name': agreement[2],
                '@inn_fop': agreement[3],
                '@pidstava_fop': agreement[4],
                '@ri_name': agreement[9],
                '@inn_ri': agreement[10],
                '@pidstava_ri': agreement[11],
                '@month_ukr_name': proto_month_ukr_name,
                '@year': proto_year,
                '@last_day_of_the_month': last_day_of_the_month,
                '@agr_sum': f"{protocol[1]:,.2f}",
                '@agrsum_handwriting_sample': protocol[2],
                '@fop_address': agreement[5],
                '@fop_iban': agreement[6],
                '@bank_account_detail_fop': agreement[7],
                '@fopname_short': agreement[8],
                '@ri_address': agreement[12],
                '@ri_iban': agreement[13],
                '@bank_account_detail_ri': agreement[14],
                '@riname_short': agreement[15]
            }


            # Замена текста в шаблоне
            self.replace_text_in_document(doc, replacements)
            self.replace_in_tables(doc.tables, replacements)
            self.formatting_text(doc)


            # Сохранение документа в память
            doc_io = BytesIO()
            doc.save(doc_io)
            doc_io.seek(0)

            # Формируем название файла
            file_name = f"{agreement[0]}_протокол_{proto_month_ukr_name}_{proto_year}.docx"

            # Отправка документа клиенту
            return await send_file(doc_io, as_attachment=True, attachment_filename=file_name)

        @self.app.route('/protocols/<int:agreement_id>/generate_act_docx/<int:protocol_id>', methods=['GET'])
        @basic_auth_required()
        async def generate_act_docx(agreement_id,protocol_id):
            # Получаем информацию по договору и протоколу
            agreement_query = """
            SELECT a.agreement_name, a.agreement_date, f.name AS fop_name, f.inn AS inn_fop, f.pidstava AS pidstava_fop, 
                   f.address AS fop_address, f.iban AS fop_iban, f.bank_account_detail AS bank_account_detail_fop, f.name_short AS fop_name_short,
                   r.name AS ri_name, r.inn AS inn_ri, r.pidstava AS pidstava_ri, r.address AS ri_address, r.iban AS ri_iban, r.bank_account_detail AS bank_account_detail_ri, r.name_short AS ri_name_short
            FROM credentials.agreements AS a
            JOIN credentials.fop_credentials AS f ON a.master_id = f.id
            JOIN credentials.ri_credentials AS r ON a.ri_id = r.id
            WHERE a.id = %s;
            """
            agreement_data = await self.local_db.execute_query(agreement_query, (agreement_id,))
            agreement = agreement_data[0]  # Мы получаем первый элемент, чтобы передать данные как строку, а не кортеж.

            protocol_query = """
            SELECT proto_date, proto_sum, proto_sum_caps
            FROM credentials.protocols
            WHERE agreement = %s and id = %s;
            """
            protocol_data = await self.local_db.execute_query(protocol_query, (agreement_id,protocol_id))
            protocol = protocol_data[0]  # Тоже получаем первый элемент из протоколов

            if not agreement_data or not protocol_data:
                return "Договор или протокол не найден", 404

            # Преобразуем данные и форматируем дату
            agreement_date_str, month_ukr_name, year, _ = self.format_date(agreement[1])
            proto_date_str, proto_month_ukr_name, proto_year, last_day_of_the_month = self.format_date(protocol[0])
            template_month = calendar.monthrange(int(protocol[0].strftime("%Y")), int(protocol[0].month))
            last_day_of_the_month = str(template_month[1])
            act_date = f"{last_day_of_the_month} {proto_month_ukr_name} {proto_year} року "
            # Загрузка шаблона
            template_path = 'static/docs/M-RI_act.docx'
            doc = Document(template_path)

            # Замена маркеров
            replacements = {
                '@agr_num': agreement[0],
                '@agr_date': agreement_date_str,
                '@proto_date': proto_date_str,
                '@fop_name': agreement[2],
                '@inn_fop': agreement[3],
                '@pidstava_fop': agreement[4],
                '@ri_name': agreement[9],
                '@inn_ri': agreement[10],
                '@pidstava_ri': agreement[11],
                '@month_ukr_name': proto_month_ukr_name,
                '@year': proto_year,
                '@last_day_of_the_month': last_day_of_the_month,
                '@agr_sum': f"{protocol[1]:,.2f}",
                '@agrsum_handwriting_sample': protocol[2],
                '@fop_address': agreement[5],
                '@fop_iban': agreement[6],
                '@bank_account_detail_fop': agreement[7],
                '@fopname_short': agreement[8],
                '@ri_address': agreement[12],
                '@ri_iban': agreement[13],
                '@bank_account_detail_ri': agreement[14],
                '@riname_short': agreement[15],
                '@today': act_date,
                '@act_nubmer': f'{protocol_id}/{agreement[0]}',
                '@act_hours': self.amount_to_time(float(protocol[1]))
            }

            # Замена текста в шаблоне
            print(self.amount_to_time(float(protocol[1])))
            self.replace_text_in_document(doc, replacements)
            self.replace_in_tables(doc.tables, replacements)
            self.formatting_text(doc)


            # Сохранение документа в память
            doc_io = BytesIO()
            doc.save(doc_io)
            doc_io.seek(0)

            # Формируем название файла
            file_name = f"{agreement[0]}_акт_{proto_month_ukr_name}_{proto_year}.docx"

            # Отправка документа клиенту
            return await send_file(doc_io, as_attachment=True, attachment_filename=file_name)

        @self.app.route('/update_agreement_state/<int:agreement_id>', methods=['POST'])
        @basic_auth_required()
        async def update_agreement_state(agreement_id):
            data = await request.json
            agreement_state = data.get('agreement_state')

            if agreement_state is None:
                return {"error": "agreement_state is required"}, 400

            query = """
            UPDATE credentials.agreements
            SET agreement_state = %s
            WHERE id = %s
            """
            await self.local_db.execute_query(query, (agreement_state, agreement_id))
            return {"message": "Agreement state updated successfully"}, 200

        @self.app.route('/agreement_detail/<int:agreement_id>', methods=['GET', 'POST'])
        @basic_auth_required()
        async def agreement_detail(agreement_id):
            # Новый SQL-запрос
            query = """
            SELECT 
                a.agreement_name, 
                a.agreement_date, 
                f.name AS master_name, 
                f.inn AS master_inn, 
                r.name AS ri_name, 
                r.inn AS ri_inn, 
                t.canton,
                a.agreement_state,
                COALESCE(MAX(term.termination_date), 'Нет данных') AS termination_date,
                GROUP_CONCAT(t.vetka ORDER BY t.vetka SEPARATOR ', ') AS vetkas
            FROM 
                credentials.agreements a
            JOIN 
                credentials.fop_credentials f ON a.master_id = f.id
            JOIN 
                credentials.ri_credentials r ON a.ri_id = r.id
            JOIN 
                credentials.fop_territory t ON a.master_id = t.master_id
            LEFT JOIN 
                credentials.agreement_termination term ON a.id = term.agreement_id
            WHERE 
                a.id = %s
            GROUP BY 
                a.agreement_name, 
                a.agreement_date, 
                f.name, 
                f.inn, 
                r.name, 
                r.inn, 
                t.canton;
            """

            # Получение данных из базы
            agreement_data = await self.local_db.execute_query(query, (agreement_id,))
            if not agreement_data:
                return "Договор не найден.", 404

            # Преобразование результата в словарь
            agreement = {
                "agreement_name": agreement_data[0][0],
                "agreement_date": agreement_data[0][1],
                "master_name": agreement_data[0][2],
                "master_inn": agreement_data[0][3],
                "ri_name": agreement_data[0][4],
                "ri_inn": agreement_data[0][5],
                "canton": agreement_data[0][6],
                "agreement_state": agreement_data[0][7],
                "termination_date": agreement_data[0][8],
                "vetkas": agreement_data[0][9],
                "agreement_id": agreement_id,
            }

            # Обработка POST-запроса для "Расторжения"
            if request.method == 'POST':
                termination_date = (await request.form).get('termination_date')
                # Логика сохранения расторжения, например, обновление базы
                term_query = """
                INSERT INTO credentials.agreement_termination(agreement_id, termination_date)
                VALUES (%s, %s)
                """
                await self.local_db.execute_query(term_query, (agreement_id, termination_date))
                return redirect(url_for('agreement_detail', agreement_id=agreement_id))

            return await render_template('agreement_detail.html', agreement=agreement)

        @self.app.route('/agreement_termination/<int:agreement_id>', methods=['GET'])
        @basic_auth_required()
        async def agreement_termination(agreement_id):
            # Запрос на получение данных о договоре
            agreement_query = """
            SELECT a.agreement_name, a.agreement_date, f.name AS fop_name, f.inn AS inn_fop, f.pidstava AS pidstava_fop, 
                   f.address AS fop_address, f.iban AS fop_iban, f.bank_account_detail AS bank_account_detail_fop, f.name_short AS fop_name_short,
                   r.name AS ri_name, r.inn AS inn_ri, r.pidstava AS pidstava_ri, r.address AS ri_address, r.iban AS ri_iban, r.bank_account_detail AS bank_account_detail_ri, r.name_short AS ri_name_short,
                   a.agreement_state, t.termination_date
            FROM credentials.agreements AS a
            JOIN credentials.fop_credentials AS f ON a.master_id = f.id
            JOIN credentials.ri_credentials AS r ON a.ri_id = r.id
            JOIN credentials.agreement_termination AS t ON a.id = t.agreement_id
            WHERE a.id = %s;
            """
            agreement_data = await self.local_db.execute_query(agreement_query, (agreement_id,))

            if not agreement_data:
                return "Договор не найден", 404

            agreement = agreement_data[0]  # Получаем первую запись для передачи данных как строку

            # Преобразуем дату договора
            agreement_date_str, month_ukr_name, year, _ = self.format_date(agreement[1])
            termination_date_str, termination_month_ukr, year, _ = self.format_date(agreement[17])

            # Загрузка шаблона расторжения договора
            template_path = 'static/docs/M-RI_termination.docx'
            doc = Document(template_path)

            # Подготовка данных для замены
            replacements = {
                '@agr_num': agreement[0],
                '@agr_date': agreement_date_str,
                '@fop_name': agreement[2],
                '@inn_fop': agreement[3],
                '@pidstava_fop': agreement[4],
                '@ri_name': agreement[9],
                '@inn_ri': agreement[10],
                '@pidstava_ri': agreement[11],
                '@fop_address': agreement[5],
                '@fop_iban': agreement[6],
                '@bank_account_detail_fop': agreement[7],
                '@fopname_short': agreement[8],
                '@ri_address': agreement[12],
                '@ri_iban': agreement[13],
                '@bank_account_detail_ri': agreement[14],
                '@riname_short': agreement[15],
                '@agreement_state': 'Активный' if agreement[16] == 1 else 'Неактивный',
                '@term_date': termination_date_str
            }

            # Замена текста в шаблоне
            self.replace_text_in_document(doc, replacements)
            self.replace_in_tables(doc.tables, replacements)
            self.formatting_text(doc)

            # Сохранение документа в память
            doc_io = BytesIO()
            doc.save(doc_io)
            doc_io.seek(0)

            # Формируем название файла
            file_name = f"{agreement[0]}_РАСТОРЖЕНИЕ_ДОГОВОРА.docx"

            # Отправка файла клиенту
            return await send_file(doc_io, as_attachment=True, attachment_filename=file_name)

        @self.app.route('/generate_contract/<int:agreement_id>', methods=['GET'])
        @basic_auth_required()
        async def generate_contract(agreement_id):
            # Запрос на получение данных о договоре
            agreement_query = """
            SELECT a.agreement_name, a.agreement_date, f.name AS fop_name, f.inn AS inn_fop, f.pidstava AS pidstava_fop, 
                   f.address AS fop_address, f.iban AS fop_iban, f.bank_account_detail AS bank_account_detail_fop, f.name_short AS fop_name_short,
                   r.name AS ri_name, r.inn AS inn_ri, r.pidstava AS pidstava_ri, r.address AS ri_address, r.iban AS ri_iban, r.bank_account_detail AS bank_account_detail_ri, r.name_short AS ri_name_short,
                   a.agreement_state
            FROM credentials.agreements AS a
            JOIN credentials.fop_credentials AS f ON a.master_id = f.id
            JOIN credentials.ri_credentials AS r ON a.ri_id = r.id
            WHERE a.id = %s;
            """
            agreement_data = await self.local_db.execute_query(agreement_query, (agreement_id,))
            agreement = agreement_data[0] # Мы получаем первый элемент, чтобы передать данные как строку, а не кортеж.

            if not agreement_data:
                return "Договор не найден", 404

            # Преобразуем данные и форматируем дату
            agreement_date_str, month_ukr_name, year, _ = self.format_date(agreement[1])

            # Загрузка шаблона договора
            template_path = 'static/docs/M-RI_agreement.docx'
            doc = Document(template_path)

            # Замена маркеров
            replacements = {
                '@agr_num': agreement[0],
                '@agr_date': agreement_date_str,
                '@fop_name': agreement[2],
                '@inn_fop': agreement[3],
                '@pidstava_fop': agreement[4],
                '@ri_name': agreement[9],
                '@inn_ri': agreement[10],
                '@pidstava_ri': agreement[11],
                '@fop_address': agreement[5],
                '@fop_iban': agreement[6],
                '@bank_account_detail_fop': agreement[7],
                '@fopname_short': agreement[8],
                '@ri_address': agreement[12],
                '@ri_iban': agreement[13],
                '@bank_account_detail_ri': agreement[14],
                '@riname_short': agreement[15],
                '@agreement_state': 'Активный' if agreement[16] == 1 else 'Неактивный'  # Активный договор или нет
            }

            # Замена текста в шаблоне
            self.replace_text_in_document(doc, replacements)
            self.replace_in_tables(doc.tables, replacements)
            self.formatting_text(doc)

            # Сохранение документа в память
            doc_io = BytesIO()
            doc.save(doc_io)
            doc_io.seek(0)

            # Формируем название файла
            file_name = f"{agreement[0]}_ДОГОВІР.docx"

            # Отправка документа клиенту
            return await send_file(doc_io, as_attachment=True, attachment_filename=file_name)

        @self.app.route('/generate_dod1/<int:agreement_id>', methods=['GET'])
        @basic_auth_required()
        async def generate_dod1(agreement_id):
            # Запросы для получения данных
            query1 = """
            SELECT 
                sr.model, 
                COUNT(DISTINCT sr.ip) AS ip_count
            FROM 
                dbsyphon.switches_report sr
            JOIN 
                credentials.fop_territory ft ON sr.vetka = ft.vetka
            JOIN 
                credentials.agreements a ON a.master_id = ft.master_id
            WHERE
                sr.switch_rank = 4
            AND
                a.id = %s
            GROUP BY 
                sr.model;
            """

            query2 = """
            SELECT 
                CONCAT(SUBSTRING_INDEX(sr.ip, '.', 3), '.0/24') AS ip_pool
            FROM 
                dbsyphon.switches_report sr
            JOIN 
                credentials.fop_territory ft ON sr.vetka = ft.vetka
            JOIN 
                credentials.agreements a ON a.master_id = ft.master_id
            WHERE
                sr.switch_rank = 4
            AND
                a.id = %s
            GROUP BY 
                ip_pool;
            """

            # Выполнение запросов и получение данных
            data_table1 = await self.local_db.execute_query(query1, (agreement_id,))
            data_table2 = await self.local_db.execute_query(query2, (agreement_id,))

            # Запрос на получение данных о договоре
            agreement_query = """
            SELECT a.agreement_name, a.agreement_date, f.name AS fop_name, f.inn AS inn_fop, f.pidstava AS pidstava_fop, 
                   f.address AS fop_address, f.iban AS fop_iban, f.bank_account_detail AS bank_account_detail_fop, f.name_short AS fop_name_short,
                   r.name AS ri_name, r.inn AS inn_ri, r.pidstava AS pidstava_ri, r.address AS ri_address, r.iban AS ri_iban, r.bank_account_detail AS bank_account_detail_ri, r.name_short AS ri_name_short,
                   a.agreement_state
            FROM credentials.agreements AS a
            JOIN credentials.fop_credentials AS f ON a.master_id = f.id
            JOIN credentials.ri_credentials AS r ON a.ri_id = r.id
            WHERE a.id = %s;
            """
            agreement_data = await self.local_db.execute_query(agreement_query, (agreement_id,))

            if not agreement_data:
                return "Договор не найден", 404

            agreement = agreement_data[0]  # Получаем данные как строку, а не кортеж.

            # Преобразуем данные и форматируем дату
            agreement_date_str, month_ukr_name, year, _ = self.format_date(agreement[1])

            # Загрузка шаблона договора
            template_path = 'static/docs/M-RI_dod1.docx'
            doc = Document(template_path)

            # Замена маркеров в шаблоне
            replacements = {
                '@agr_num': agreement[0],
                '@agr_date': agreement_date_str,
                '@fop_name': agreement[2],
                '@inn_fop': agreement[3],
                '@pidstava_fop': agreement[4],
                '@ri_name': agreement[9],
                '@inn_ri': agreement[10],
                '@pidstava_ri': agreement[11],
                '@fop_address': agreement[5],
                '@fop_iban': agreement[6],
                '@bank_account_detail_fop': agreement[7],
                '@fopname_short': agreement[8],
                '@ri_address': agreement[12],
                '@ri_iban': agreement[13],
                '@bank_account_detail_ri': agreement[14],
                '@riname_short': agreement[15],
                '@agreement_state': 'Активний' if agreement[16] == 1 else 'Неактивний'  # Активный договор или нет
            }

            # Замена текста в шаблоне
            self.replace_text_in_document(doc, replacements)
            self.replace_in_tables(doc.tables, replacements)
            self.formatting_text(doc)

            # Создание таблицы @table1
            table1 = self.create_table(doc, data_table1,
                                       ['Найменування (модель) технічних засобів електронних комунікацій', 'Кількість'])

            # Создание таблицы @table2
            table2 = self.create_table(doc, data_table2,
                                       ['Діапазон ІР адрес технічних засобів електронних комунікацій'])

            # Замена маркеров @table1 и @table2 на таблицы
            self.replace_table_in_document(doc, '@table1', table1)
            self.replace_table_in_document(doc, '@table2', table2)

            # Сохранение документа в память
            doc_io = BytesIO()
            doc.save(doc_io)
            doc_io.seek(0)

            # Формируем название файла
            file_name = f"{agreement[0]}_Додаток1.docx"

            # Отправка документа клиенту
            return await send_file(doc_io, as_attachment=True, attachment_filename=file_name)


        @self.app.route('/protocols/<int:agreement_id>', methods=['GET', 'POST'])
        @basic_auth_required()
        async def protocols(agreement_id):
            if request.method == 'POST':
                # Получение данных из формы
                proto_date = (await request.form)['proto_date']
                proto_sum = (await request.form)['proto_sum']
                # proto_sum_caps = (await request.form)['proto_sum_caps']
                proto_sum_caps = self.convert_to_currency_words(float(proto_sum))

                # Вставка данных протокола в базу данных
                insert_query = """
                INSERT INTO credentials.protocols (agreement, proto_date, proto_sum, proto_sum_caps)
                VALUES (%s, %s, %s, %s);
                """
                await self.local_db.execute_query(insert_query, (agreement_id, proto_date, proto_sum, proto_sum_caps))

                # Перезагрузка страницы после добавления данных
                return redirect(url_for('protocols', agreement_id=agreement_id))

            # Запрос существующих протоколов с proto_state = 1
            protocols_query = """
            SELECT p.proto_date, p.proto_sum, p.proto_sum_caps, p.proto_state, p.id
            FROM credentials.protocols AS p
            WHERE p.agreement = %s AND p.proto_state = 1;
            """
            protocols_data = await self.local_db.execute_query(protocols_query, (agreement_id,))

            # Проверка на пустой результат запроса
            if protocols_data is None:
                protocols_data = []

            # Преобразование кортежей в словари
            protocols = [
                {
                    'proto_date': protocol[0],
                    'proto_sum': protocol[1],
                    'proto_sum_caps': protocol[2],
                    'proto_state': protocol[3],
                    'id': protocol[4]
                }
                for protocol in protocols_data
            ]

            # Запрос данных о договоре
            agreement_query = """
            SELECT a.agreement_name, f.name AS master_name, r.name AS ri_name, a.id, f.email
            FROM credentials.agreements AS a
            JOIN credentials.fop_credentials AS f ON a.master_id = f.id
            JOIN credentials.ri_credentials AS r ON a.ri_id = r.id
            WHERE a.id = %s;
            """
            agreement = await self.local_db.execute_query(agreement_query, (agreement_id,))

            if not agreement:
                return "Договор не найден", 404  # Или можно сделать редирект на другую страницу

            # Отображение страницы протоколов
            return await render_template('protocols.html', protocols=protocols, agreement=agreement[0],
                                         agreement_id=agreement_id)

        @self.app.route('/protocols/<int:agreement_id>/delete/<int:protocol_id>', methods=['POST'])
        @basic_auth_required()
        async def delete_protocol(agreement_id, protocol_id):
            update_query = """
            UPDATE credentials.protocols 
            SET proto_state = 0 
            WHERE id = %s
            """
            await self.local_db.execute_query(update_query, (protocol_id,))

            # Перезагрузка страницы после удаления
            return redirect(url_for('protocols', agreement_id=agreement_id))

        @self.app.route('/agreements', methods=['GET'])
        @basic_auth_required()
        async def agreements():

            # Получаем выбранный фильтр из параметров запроса
            selected_engineer = request.args.get('engineer_filter', '')
            selected_canton = request.args.get('canton_filter', '')
            selected_state = request.args.get('state_filter', 'all')
            # SQL-запрос для получения всех инженеров
            engineers_query = "SELECT DISTINCT name FROM credentials.ri_credentials;"
            engineers_data = await self.local_db.execute_query(engineers_query)
            engineers = [row[0] for row in engineers_data]
            # Получаем список всех округов
            cantons_query = "SELECT DISTINCT canton FROM credentials.fop_territory;"
            cantons_data = await self.local_db.execute_query(cantons_query)
            cantons = [row[0] for row in cantons_data]

            # SQL-запрос для получения информации о всех договорах
            query = """
                                SELECT 
                a.id, 
                a.agreement_name, 
                fc.name AS master_name, 
                rc.name AS engineer_name, 
                a.agreement_state, 
                ter.canton
            FROM 
                credentials.agreements a
            JOIN 
                credentials.fop_credentials fc ON a.master_id = fc.id
            JOIN 
                credentials.ri_credentials rc ON a.ri_id = rc.id
            JOIN 
                (
                    SELECT DISTINCT master_id, canton
                    FROM credentials.fop_territory where transition is null
                ) ter ON ter.master_id = a.master_id
            """

            # Условия для фильтров
            filters = []
            params = []

            if selected_engineer:
                filters.append("rc.name = %s")
                params.append(selected_engineer)

            if selected_canton:
                filters.append("ter.canton = %s")
                params.append(selected_canton)

            if selected_state == 'active':
                filters.append("a.agreement_state = 1")
            elif selected_state == 'inactive':
                filters.append("a.agreement_state = 0")

            # Добавляем условия в запрос
            if filters:
                query += " WHERE " + " AND ".join(filters)

            # Выполняем запрос
            agreements_data = await self.local_db.execute_query(query, tuple(params))


            # # Выполняем запрос и получаем данные по договорам
            # agreements_data = await self.local_db.execute_query(query)

            all_years = set()
            agreements_list = []

            # Преобразуем каждый кортеж в словарь
            for agreement in agreements_data:
                agreement_dict = {
                    'id': agreement[0],
                    'agreement_name': agreement[1],
                    'master_name': agreement[2],
                    'engineer_name': agreement[3],
                    'agreement_state': agreement[4],
                    'canton': agreement[5],
                    'protocols_by_year': {}
                }

                # Для каждого договора получаем данные по протоколам
                proto_query = """
                SELECT YEAR(proto_date) AS proto_year, MONTH(proto_date) AS proto_month
                FROM credentials.protocols
                WHERE proto_state = 1 AND agreement = %s
                ORDER BY proto_year, proto_month;
                """
                proto_data = await self.local_db.execute_query(proto_query, (agreement[0],))

                # Сортируем протоколы по годам
                for row in proto_data:
                    year, month = row
                    if year not in agreement_dict['protocols_by_year']:
                        agreement_dict['protocols_by_year'][year] = []
                    agreement_dict['protocols_by_year'][year].append(month)
                    all_years.add(year)

                agreements_list.append(agreement_dict)

            # Передаем все данные в шаблон, включая все уникальные годы

            return await render_template(
                'agreements.html',
                agreements=agreements_list,
                all_years=sorted(all_years),
                engineers=engineers,
                cantons=cantons,
                selected_engineer=selected_engineer,
                selected_canton=selected_canton,
                selected_state=selected_state
            )

        @self.app.route('/fop-form')
        @basic_auth_required()
        async def fop_form():
            # Отображение формы для внесения информации о ФОП
            return await render_template('fop_form.html')

        @self.app.route('/submit-fop', methods=['POST'])
        async def submit_fop():
            # Используем await для получения данных формы
            form_data = await request.form

            # Получаем значения полей из формы
            position = form_data.get('position')
            name = form_data.get('name')
            inn = form_data.get('inn')
            pidstava = form_data.get('pidstava')
            address = form_data.get('address')
            iban = form_data.get('iban')
            bank_account_detail = form_data.get('bank_account_detail')
            name_short = form_data.get('name_short')
            email = form_data.get('email')
            canton = form_data.get('canton')
            vetkas = form_data.getlist('vetka[]')  # Получаем список веток

            # Выбор таблицы в зависимости от позиции
            if position == 'Мастер':
                table = 'credentials.fop_credentials'
            elif position == 'Инженер':
                table = 'credentials.ri_credentials'
            else:
                return jsonify({"error": "Неверная позиция"}), 400

            # SQL-запрос для вставки данных в fop_credentials
            insert_query = f"""
            INSERT INTO {table} (name, inn, pidstava, address, iban, bank_account_detail, name_short, email)
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
            """

            # Вставка данных в базу и получение id новой записи
            try:
                # Выполняем вставку данных
                await self.local_db.execute_query(insert_query, (
                name, inn, pidstava, address, iban, bank_account_detail, name_short, email))

                # Получаем id последней вставленной записи
                result = await self.local_db.execute_query("select max(id) from credentials.fop_credentials")
                master_id = result[0][0]  # Получаем id новой записи
                print(master_id)

                # Вставка данных в таблицу fop_territory
                if vetkas:
                    insert_territory_query = """
                    INSERT INTO credentials.fop_territory (master_id, canton, vetka)
                    VALUES (%s, %s, %s)
                    """
                    for vetka in vetkas:
                        print(vetka)
                        await self.local_db.execute_query(insert_territory_query, (master_id, canton, int(vetka)))

                return jsonify({"message": "Данные успешно добавлены"}), 200
            except Exception as e:
                return jsonify({"error": str(e)}), 500

        @self.app.route('/check-inn')
        async def check_inn():
            inn = request.args.get('inn')
            position = request.args.get('position')

            if not inn or not position:
                return jsonify({"error": "Отсутствуют обязательные параметры"}), 400

            if position == 'Мастер':
                table = 'credentials.fop_credentials'
            elif position == 'Инженер':
                table = 'credentials.ri_credentials'
            else:
                return jsonify({"error": "Неверная позиция"}), 400

            try:
                quoted_table = quote_qualified_identifier(table)
                query = f"SELECT EXISTS(SELECT 1 FROM {quoted_table} WHERE inn = %s)"
                result = await self.local_db.execute_query(query, (inn,))
                exists = result[0][0]  # Результат запроса (True или False)

                return jsonify({"exists": exists}), 200
            except Exception as e:
                return jsonify({"error": str(e)}), 500

        @self.app.route('/')
        @self.app.route('/index')
        @basic_auth_required()
        async def index():
            return await render_template('index.html')

        @self.app.route('/api-test')
        @basic_auth_required()
        async def api_test():
            remote_methods = [
                {
                    "name": "fetch-logs-and-store",
                    "method": "GET",
                    "endpoint": "fetch_logs_and_store",
                    "route": "/fetch-logs-and-store",
                    "remote_source": "command_logs.logs, mrtg.switches",
                    "local_target": "dbsyphon.ntst_logs",
                },
                {
                    "name": "fetch-and-store",
                    "method": "GET",
                    "endpoint": "fetch_and_store",
                    "route": "/fetch-and-store",
                    "remote_source": "pinger.HostLogs",
                    "local_target": "dbsyphon.ntst_pinger_hosts_log",
                },
                {
                    "name": "bdcom_list",
                    "method": "POST",
                    "endpoint": "bdcom_list",
                    "route": "/bdcom_list",
                    "remote_source": "mrtg.switches",
                    "local_target": "dbsyphon.bdcom_list",
                },
                {
                    "name": "sync-switches-report",
                    "method": "GET",
                    "endpoint": "sync_switches_report",
                    "route": "/sync-switches-report",
                    "remote_source": "mrtg.switches",
                    "local_target": "dbsyphon.switches_report",
                },
            ]

            base_url = request.url_root.rstrip("/")
            for item in remote_methods:
                item["url"] = url_for(item["endpoint"])
                item["absolute_url"] = f"{base_url}{item['url']}"

            return await render_template("api_test.html", remote_methods=remote_methods)

        @self.app.route('/device_report')
        @basic_auth_required()
        async def device_report():
            devices_query = "SELECT id, description FROM devices WHERE device_type = 'power_control';"
            devices = await self.local_db.execute_query(devices_query)
            return await render_template('dev-report.html', devices=devices)

        @self.app.route('/generate-report', methods=['GET'])
        @basic_auth_required()
        async def generate_report():
            start_date = request.args.get('start_date')
            end_date = request.args.get('end_date')

            # Получаем все связки устройств из таблицы device_relations
            relations_query = """
            SELECT d1.id AS power_id, d1.description AS power_desc, d2.id AS generator_id, d2.description AS generator_desc
            FROM device_relations AS r
            JOIN devices AS d1 ON r.power_control_id = d1.id
            JOIN devices AS d2 ON r.generator_control_id = d2.id;
            """
            relations_data = await self.local_db.execute_query(relations_query)

            table_data = []

            for relation in relations_data:
                power_id, power_desc, generator_id, generator_desc = relation

                # Запрос данных для power_control
                power_downtime_query = """
                SELECT SUM(downtime)
                FROM ntst_pinger_hosts_log
                WHERE ip = (SELECT ip_address FROM devices WHERE id = %s)
                AND start BETWEEN %s AND %s;
                """
                power_downtime_data = await self.local_db.execute_query(power_downtime_query,
                                                                        (power_id, start_date, end_date))
                total_power_downtime = power_downtime_data[0][0] if power_downtime_data[0][0] else 0

                # Запрос данных для generator_control
                generator_downtime_query = """
                SELECT SUM(downtime)
                FROM ntst_pinger_hosts_log
                WHERE ip = (SELECT ip_address FROM devices WHERE id = %s)
                AND start BETWEEN %s AND %s;
                """
                generator_downtime_data = await self.local_db.execute_query(generator_downtime_query,
                                                                            (generator_id, start_date, end_date))
                total_generator_downtime = generator_downtime_data[0][0] if generator_downtime_data[0][0] else 0

                # Рассчет uptime генератора во время простоя power_control
                generator_uptime_during_power_downtime = max(0,
                                                             (total_power_downtime - total_generator_downtime) / 3600)

                # Добавление строки в таблицу с округлением до 2 знаков после запятой
                table_data.append({
                    'description': f"{power_desc}",
                    'power_downtime_hours': round(total_power_downtime / 3600, 2),  # в часах
                    'generator_downtime_hours': round(total_generator_downtime / 3600, 2),  # в часах
                    'generator_uptime_during_power_downtime_hours': round(generator_uptime_during_power_downtime, 2)
                })

            # Дополнительный запрос для суммарного времени простоя всех устройств power_control
            total_downtime_query = """
            SELECT d.description, SUM(l.downtime) / 3600 AS total_downtime_hours
            FROM devices AS d
            LEFT JOIN ntst_pinger_hosts_log AS l ON d.ip_address = l.ip
            WHERE d.device_type = 'power_control'
            AND l.start BETWEEN %s AND %s
            GROUP BY d.description;
            """
            total_downtime_data = await self.local_db.execute_query(total_downtime_query, (start_date, end_date))

            # Передача данных в шаблон
            return await render_template('report.html',
                                         table_data=table_data,
                                         total_downtime_data=total_downtime_data,
                                         start_date=start_date,
                                         end_date=end_date
                                         )

        # @self.app.route('/')
        # async def index():
        #     return jsonify({"message": "Добро пожаловать в Quart приложение!"})

        @self.app.route('/fetch-logs-and-store', methods=['GET'])
        async def fetch_logs_and_store():
            """
            Подключается к удаленной базе данных, извлекает данные из таблиц command_logs.logs и mrtg.switches
            для определенных кантонов и сохраняет их в локальной базе данных dbsyphon.ntst_logs,
            пропуская записи с уже существующими идентификаторами.
            """
            # Подключение к удаленной базе данных
            try:
                await self.remote_db.connect()
            except Exception as e:
                return jsonify({"error": "Удаленная база данных недоступна, попробуйте позже."}), 503

            # Запрос данных из удаленной базы данных
            remote_query = """
            SELECT log.id, log.datetime, log.ip, sw.canton, sw.model, sw.rank
            FROM command_logs.logs log
            JOIN mrtg.switches sw ON log.ip = sw.ip
            WHERE sw.canton IN (
                'Минский', 'Оболонский', 'Голосеевский', 'Виноградарский', 
                'Лукьяновский', 'Святошинский', 'Бощаговский', 'Теремковский'
            )
            """
            remote_data = await self.remote_db.execute_query(remote_query)

            # Закрытие соединения с удаленной базой данных
            await self.remote_db.close()

            if not remote_data:
                return jsonify({
                    "message": "Данные успешно получены, но новых данных нет.",
                    "new_data_count": 0
                }), 200

            # Вставка данных в локальную базу данных
            insert_query = """
            INSERT INTO dbsyphon.ntst_logs (id, log_date, ip, canton, model, sw_rank)
            VALUES (%s, %s, %s, %s, %s, %s)
            """
            check_query = "SELECT COUNT(*) FROM dbsyphon.ntst_logs WHERE id = %s"

            new_data_count = 0
            for record in remote_data:
                # Проверка наличия записи в локальной базе данных
                existing_count = await self.local_db.execute_query(check_query, (record[0],))

                # Если запись не существует, выполняем вставку
                if existing_count[0][0] == 0:
                    await self.local_db.execute_query(
                        insert_query,
                        (record[0], record[1], record[2], record[3], record[4], record[5])
                    )
                    new_data_count += 1

            return jsonify({
                "message": "Данные успешно получены и сохранены в локальной базе данных.",
                "new_data_count": new_data_count
            })
        @self.app.route('/fetch-and-store', methods=['GET'])
        async def fetch_and_store():
            """
            Проверяет доступность удаленной базы данных, получает список IP-адресов из локальной базы,
            извлекает данные из удаленной базы для этих IP-адресов и сохраняет их в локальной базе данных.
            Возвращает количество новых записей, которые были добавлены.
            """
            # Проверка доступности удаленной базы данных
            try:
                await self.remote_db.connect()
            except Exception as e:
                return jsonify({"error": "Удаленная база данных недоступна, попробуйте позже."}), 503

            # Получение списка IP-адресов из локальной базы данных
            ip_query = "SELECT ip_address FROM dbsyphon.devices"
            ip_addresses = await self.local_db.execute_query(ip_query)

            if not ip_addresses:
                await self.remote_db.close()
                return jsonify({"error": "Не удалось получить IP-адреса из локальной базы данных."}), 500

            # Формируем список IP-адресов для SQL запроса
            ip_list = [ip[0] for ip in ip_addresses]

            # Получение всех ID из локальной базы данных
            local_ids_query = "SELECT id FROM ntst_pinger_hosts_log;"
            local_ids = await self.local_db.execute_query(local_ids_query)
            local_ids_list = [id[0] for id in local_ids]

            ip_placeholders = build_placeholders(ip_list)
            remote_params = list(ip_list)
            id_filter = ""
            if local_ids_list:
                id_placeholders = build_placeholders(local_ids_list)
                id_filter = f"AND hl.id NOT IN ({id_placeholders})"
                remote_params.extend(local_ids_list)

            # Запрос данных из удаленной базы данных для полученных IP-адресов
            remote_query = f"""
            SELECT hl.ip, hl.start, hl.stop, hl.id
            FROM pinger.HostLogs AS hl
            WHERE hl.ip IN ({ip_placeholders})
            {id_filter};
            """
            remote_data = await self.remote_db.execute_query(remote_query, tuple(remote_params))

            # Закрытие соединения с удаленной базой данных
            await self.remote_db.close()

            if not remote_data:
                return jsonify({
                    "message": "Данные успешно получены, но новых данных нет.",
                    "new_data_count": 0
                }), 200

            # Подсчет количества новых данных
            new_data_count = len(remote_data)

            # Вставка данных в локальную базу данных
            insert_query = """
            INSERT IGNORE INTO dbsyphon.ntst_pinger_hosts_log (id, ip, start, stop, downtime)
            VALUES (%s, %s, %s, %s, TIMESTAMPDIFF(SECOND, %s, %s))
            """

            for record in remote_data:
                await self.local_db.execute_query(
                    insert_query,
                    (record[3], record[0], record[1], record[2], record[1], record[2])
                )

            return jsonify({
                "message": "Данные успешно получены и сохранены в локальной базе данных.",
                "new_data_count": new_data_count
            })

        @self.app.route('/bdcom_list', methods=['POST'])
        async def bdcom_list():
            try:
                # Очистка таблицы
                await self.local_db.execute_query("TRUNCATE TABLE dbsyphon.bdcom_list")

                # Запрос к remote_db
                remote_query = """
                    SELECT id AS ntst_id, ip, login, PASSWORD AS passwd
                    FROM mrtg.switches
                    WHERE canton IN (
                        'Голосеевский', 'Минский', 'Лукьяновский', 'Оболонский',
                        'Виноградарский', 'Борщаговский', 'Теремковский', 'Святошинский'
                    ) AND model LIKE 'BDCOM%';
                """
                results = await self.remote_db.execute_query(remote_query)

                # Вставка в локальную таблицу
                insert_query = """
                    INSERT INTO dbsyphon.bdcom_list (ntst_id, ip, login, passwd)
                    VALUES (%s, %s, %s, %s);
                """
                if results:
                    for row in results:
                        await self.local_db.execute_query(insert_query, row)

                return jsonify({"status": "success", "inserted": len(results)}), 200

            except Exception as e:
                import traceback
                traceback.print_exc()
                return jsonify({"status": "error", "message": str(e)}), 500

        @self.app.route('/bdcom_list/export', methods=['GET'])
        @basic_auth_required()
        async def export_bdcom_list():
            try:
                # Получаем все записи из bdcom_list
                query = "SELECT id, ntst_id, ip, login, passwd FROM dbsyphon.bdcom_list;"
                rows = await self.local_db.execute_query(query)

                # Преобразуем в список словарей
                result = [
                    {
                        "id": row[0],
                        "ntst_id": row[1],
                        "ip": row[2],
                        "login": row[3],
                        "passwd": row[4]
                    }
                    for row in rows
                ]

                return jsonify(result), 200

            except Exception as e:
                import traceback
                traceback.print_exc()
                return jsonify({"status": "error", "message": str(e)}), 500

        @self.app.route('/sync-switches-report', methods=['GET'])
        async def sync_switches_report():
            """
            Синхронизирует данные из удаленной таблицы `switches` с локальной базой данных.
            Проверяет доступность удаленной базы данных и использует `fetch_info` для отслеживания последней синхронизации.
            """

            # Шаг 1: Проверка доступности удаленной базы данных
            try:
                await self.remote_db.connect()
            except Exception as e:
                return jsonify({"error": "Удаленная база данных недоступна, попробуйте позже."}), 503

            # Шаг 2: Получение времени последней модификации из локальной базы данных
            modification_time_query = """
            SELECT modification_time 
            FROM dbsyphon.fetch_info 
            WHERE db = 'dbsyphon' AND db_table = 'switches_report';
            """
            modification_time_result = await self.local_db.execute_query(modification_time_query)

            # Получаем текущую дату
            current_date = date.today()

            if not modification_time_result:
                # Вариант 1: Пустой ответ - запись отсутствует, значит, синхронизация еще не выполнялась
                # Вставляем запись в `fetch_info` с сегодняшней датой
                insert_fetch_info_query = """
                INSERT INTO dbsyphon.fetch_info (db, db_table, modification_time)
                VALUES ('dbsyphon', 'switches_report', %s);
                """
                await self.local_db.execute_query(insert_fetch_info_query, (current_date,))

                # Выполняем запрос к удаленной базе для получения всех нужных данных
                remote_query = """
                SELECT sw.canton, sw.model, sw.ip, sw.rank, sw.vetka
                FROM mrtg.switches sw
                WHERE sw.canton IN ('Минский', 'Оболонский', 'Голосеевский', 'Виноградарский', 
                                    'Лукьяновский', 'Святошинский', 'Борщаговский', 'Теремковский') 
                    AND sw.model NOT IN ('Контроль питания, ранг 3', 
                                         'Контроль питания, ранг 2', 
                                         'Контроль питания, генератор', 
                                         'DGS-1100-06/ME R3', 
                                         'Датчик дыма');
                """
                remote_data = await self.remote_db.execute_query(remote_query)

                # Вставка данных в локальную таблицу `switches_report`
                insert_switches_report_query = """
                INSERT INTO dbsyphon.switches_report (canton, model, ip, switch_rank, vetka)
                VALUES (%s, %s, %s, %s, %s);
                """
                for record in remote_data:
                    await self.local_db.execute_query(insert_switches_report_query, record)

                await self.remote_db.close()
                return jsonify({"message": "Данные успешно получены и добавлены в switches_report."}), 200

            # Обработка случая, когда результат есть
            last_modification_date = modification_time_result[0][0]

            if last_modification_date == current_date:
                # Вариант 2: Синхронизация уже выполнена сегодня
                await self.remote_db.close()
                return jsonify({"message": "Синхронизация уже проводилась сегодня."}), 200

            else:
                # Вариант 3: Дата модификации отличается от сегодняшней - создаем архивную таблицу
                # Создание таблицы history.switches_report_<сегодняшняя_дата>
                archive_table_name = f"history.switches_report_{current_date.strftime('%Y_%m_%d')}"
                quoted_archive_table = quote_qualified_identifier(archive_table_name)
                create_archive_table_query = f"""
                CREATE TABLE {quoted_archive_table} (
                    id INT AUTO_INCREMENT PRIMARY KEY,
                    canton VARCHAR(100),
                    model VARCHAR(100),
                    ip VARCHAR(15),
                    switch_rank TINYINT UNSIGNED,
                    vetka INT
                );
                """
                await self.local_db.execute_query(create_archive_table_query)

                # Копируем данные из `switches_report` в архивную таблицу
                copy_to_archive_query = f"""
                INSERT INTO {quoted_archive_table} (canton, model, ip, switch_rank, vetka)
                SELECT canton, model, ip, switch_rank, vetka
                FROM dbsyphon.switches_report;
                """
                await self.local_db.execute_query(copy_to_archive_query)

                # Очищаем данные из `switches_report`
                clear_switches_report_query = "DELETE FROM dbsyphon.switches_report;"
                await self.local_db.execute_query(clear_switches_report_query)

                # Обновляем запись в `fetch_info` с новой датой синхронизации
                update_fetch_info_query = """
                UPDATE dbsyphon.fetch_info 
                SET modification_time = %s 
                WHERE db = 'dbsyphon' AND db_table = 'switches_report';
                """
                await self.local_db.execute_query(update_fetch_info_query, (current_date,))
                insert_switches_report_query = """
                INSERT INTO dbsyphon.switches_report (canton, model, ip, switch_rank, vetka)
                VALUES (%s, %s, %s, %s, %s);
                """

                # Выполняем запрос к удаленной базе для получения новых данных
                remote_query = """
                SELECT sw.canton, sw.model, sw.ip, sw.rank, sw.vetka
                FROM mrtg.switches sw
                WHERE sw.canton IN ('Минский', 'Оболонский', 'Голосеевский', 'Виноградарский', 
                                    'Лукьяновский', 'Святошинский', 'Борщаговский', 'Теремковский') 
                    AND sw.model NOT IN ('Контроль питания, ранг 3', 
                                         'Контроль питания, ранг 2', 
                                         'Контроль питания, генератор', 
                                         'DGS-1100-06/ME R3', 
                                         'Датчик дыма');
                """
                remote_data = await self.remote_db.execute_query(remote_query)

                # Вставка данных в `switches_report` после очистки
                for record in remote_data:
                    print(record)
                    await self.local_db.execute_query(insert_switches_report_query, record)

                await self.remote_db.close()
                return jsonify({"message": "Данные успешно обновлены в switches_report."}), 200

    def run(self):
        # Запуск приложения на Quart
        self.app.run(debug=True)

app = MyApp()
asgi_app = app.app



if __name__ == '__main__':
    app = MyApp()
    app.run()
    # asgi_app.run()
